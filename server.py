import io
import os
import sys
import json
import uuid
import gc
import re
import requests
import smtplib
import hashlib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from datetime import datetime, timedelta
from typing import List, Optional, Dict, Tuple

import numpy as np
import pandas as pd
import pdfplumber
import faiss
import anthropic

from fastapi import FastAPI, UploadFile, File, HTTPException, BackgroundTasks, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, HTMLResponse
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel
from docx import Document
from docx.shared import RGBColor
from openai import OpenAI
from urllib.parse import urlparse, urljoin
from bs4 import BeautifulSoup

# ── Inline Auth & Database ────────────────────────────────────────────────────
AUTH_ENABLED = False
try:
    from sqlalchemy import (create_engine, Column, Integer, String, Boolean,
                            DateTime, Float, ForeignKey)
    from sqlalchemy.ext.declarative import declarative_base
    from sqlalchemy.orm import sessionmaker, relationship
    from jose import JWTError, jwt

    _DB_URL = os.environ.get("DATABASE_URL", "sqlite:///./cybesure.db")
    if _DB_URL.startswith("postgres://"):
        _DB_URL = _DB_URL.replace("postgres://", "postgresql://", 1)
    _connect_args = {"check_same_thread": False} if "sqlite" in _DB_URL else {}
    _engine = create_engine(_DB_URL, connect_args=_connect_args)
    _SessionLocal = sessionmaker(autocommit=False, autoflush=False, bind=_engine)
    _Base = declarative_base()

    SUBSCRIPTION_CONFIG = {
        "demo":      {"name":"Demo",       "limit":3,   "price_gbp":0,     "q_limit":10},
        "starter":   {"name":"Starter",    "limit":50,  "price_gbp":5000,  "q_limit":None},
        "business":  {"name":"Business",   "limit":100, "price_gbp":8500,  "q_limit":None},
        "corporate": {"name":"Corporate",  "limit":200, "price_gbp":15000, "q_limit":None},
        "enterprise":{"name":"Enterprise", "limit":500, "price_gbp":25000, "q_limit":None},
        "test":      {"name":"Test",       "limit":999, "price_gbp":0,     "q_limit":None},
    }

    TOPUP_CONFIG = {
        "five":   {"qty":5,  "price_gbp":750,  "label":"5 questionnaires"},
        "ten":    {"qty":10, "price_gbp":1500, "label":"10 questionnaires"},
    }

    # WooCommerce config
    WC_SITE_URL    = os.environ.get("WC_SITE_URL",    "https://cybesure.com")
    WC_CONSUMER_KEY    = os.environ.get("WC_CONSUMER_KEY",    "")
    WC_CONSUMER_SECRET = os.environ.get("WC_CONSUMER_SECRET", "")

    class Organisation(_Base):
        __tablename__ = "organisations"
        id                  = Column(Integer, primary_key=True, index=True)
        name                = Column(String(200), nullable=False)
        slug                = Column(String(100), unique=True, index=True)
        contact_email       = Column(String(200), nullable=False)
        created_at          = Column(DateTime, default=datetime.utcnow)
        tier                = Column(String(50), default="starter")
        status              = Column(String(50), default="trial")
        questionnaire_limit = Column(Integer, default=50)
        questionnaires_used = Column(Integer, default=0)
        topup_credits       = Column(Integer, default=0)
        subscription_start  = Column(DateTime, default=datetime.utcnow)
        subscription_end    = Column(DateTime, default=lambda: datetime.utcnow() + timedelta(days=14))
        users               = relationship("User", back_populates="organisation")

        @property
        def total_limit(self): return self.questionnaire_limit + self.topup_credits
        @property
        def remaining(self): return max(0, self.total_limit - self.questionnaires_used)
        @property
        def is_active(self):
            return self.status in ("active","trial") and self.subscription_end > datetime.utcnow()

    class User(_Base):
        __tablename__ = "users"
        id              = Column(Integer, primary_key=True, index=True)
        org_id          = Column(Integer, ForeignKey("organisations.id"), nullable=False)
        email           = Column(String(200), unique=True, index=True, nullable=False)
        full_name       = Column(String(200), nullable=False)
        hashed_password = Column(String(200), nullable=False)
        role            = Column(String(50), default="contributor")
        is_active       = Column(Boolean, default=True)
        created_at      = Column(DateTime, default=datetime.utcnow)
        last_login      = Column(DateTime, nullable=True)
        organisation    = relationship("Organisation", back_populates="users")

    class QuestionnaireRun(_Base):
        __tablename__ = "questionnaire_runs"
        id              = Column(Integer, primary_key=True, index=True)
        org_id          = Column(Integer, ForeignKey("organisations.id"), nullable=False)
        user_id         = Column(Integer, ForeignKey("users.id"), nullable=False)
        filename        = Column(String(500), nullable=True)
        question_count  = Column(Integer, default=0)
        status          = Column(String(50), default="processing")
        created_at      = Column(DateTime, default=datetime.utcnow)
        completed_at    = Column(DateTime, nullable=True)
        industry_sector = Column(String(200), nullable=True)
        processing_secs = Column(Float, default=0)   # seconds to process
        avg_complexity  = Column(String(50), default="average")  # simple/average/complex

    _Base.metadata.create_all(bind=_engine)

    _JWT_SECRET = os.environ.get("JWT_SECRET", "cybesure-secureanswer-2025")

    def hash_password(p: str) -> str:
        """Secure password hashing using SHA256 + salt — no bcrypt size limits."""
        salt = "cybesure-salt-2025"
        return hashlib.sha256(f"{salt}{p}{salt}".encode()).hexdigest()

    def verify_password(plain: str, hashed: str) -> bool:
        return hash_password(plain) == hashed

    def make_token(data: dict) -> str:
        d = data.copy()
        d["exp"] = datetime.utcnow() + timedelta(hours=24)
        return jwt.encode(d, _JWT_SECRET, algorithm="HS256")

    def get_woocommerce_subscription(email: str) -> Optional[str]:
        """
        Check WooCommerce subscriptions for a given email.
        Returns the tier name if active subscription found, None otherwise.
        Maps WooCommerce product names to subscription tiers.
        """
        if not WC_CONSUMER_KEY or not WC_CONSUMER_SECRET:
            return None
        try:
            # WooCommerce REST API - get subscriptions by email
            resp = requests.get(
                f"{WC_SITE_URL}/wp-json/wc/v3/subscriptions",
                params={"search": email, "status": "active"},
                auth=(WC_CONSUMER_KEY, WC_CONSUMER_SECRET),
                timeout=10
            )
            if resp.status_code != 200:
                return None
            subs = resp.json()
            # Map product names to tiers
            tier_map = {
                "starter": "starter", "business": "business",
                "corporate": "corporate", "enterprise": "enterprise",
                "secureanswer starter": "starter",
                "secureanswer business": "business",
                "secureanswer corporate": "corporate",
                "secureanswer enterprise": "enterprise",
            }
            for sub in subs:
                for item in sub.get("line_items", []):
                    product_name = item.get("name", "").lower()
                    for key, tier in tier_map.items():
                        if key in product_name:
                            return tier
        except Exception as e:
            print(f"WooCommerce check error: {e}")
        return None

    def get_db():
        db = _SessionLocal()
        try: yield db
        finally: db.close()

    def decode_token(token: str):
        """Decode a JWT token and return the user."""
        try:
            payload = jwt.decode(token, _JWT_SECRET, algorithms=["HS256"])
            uid = payload.get("sub")
            if not uid:
                raise HTTPException(401, "Invalid token")
        except JWTError:
            raise HTTPException(401, "Invalid token")
        db = _SessionLocal()
        try:
            user = db.query(User).filter(User.id==int(uid), User.is_active==True).first()
        finally:
            db.close()
        if not user:
            raise HTTPException(401, "User not found")
        return user

    def get_auth_user(token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
        return decode_token(token)

    AUTH_ENABLED = True
    print("Auth system initialised successfully")

except Exception as _auth_err:
    import traceback as _tb
    print(f"Auth not available: {_auth_err}")
    print(_tb.format_exc())
    async def get_db(): yield None
    def get_auth_user(): return None
    def decode_token(token): raise HTTPException(503, "Auth not available")
    def hash_password(p):
        salt = "cybesure-salt-2025"
        return hashlib.sha256(f"{salt}{p}{salt}".encode()).hexdigest()
    def verify_password(p, h): return hash_password(p) == h
    def make_token(d): return "no-auth"
    def check_subscription(org): pass
    def get_woocommerce_subscription(email): return None
    WC_CONSUMER_KEY = ""
    WC_CONSUMER_SECRET = ""
    WC_SITE_URL = ""
    SUBSCRIPTION_CONFIG = {
        "demo":{"name":"Demo","limit":3,"price_gbp":0,"q_limit":10},
        "starter":{"name":"Starter","limit":50,"price_gbp":5000,"q_limit":None},
        "business":{"name":"Business","limit":100,"price_gbp":8500,"q_limit":None},
        "corporate":{"name":"Corporate","limit":200,"price_gbp":15000,"q_limit":None},
        "enterprise":{"name":"Enterprise","limit":500,"price_gbp":25000,"q_limit":None},
        "test":{"name":"Test","limit":999,"price_gbp":0,"q_limit":None},
    }
    TOPUP_CONFIG = {
        "five":{"qty":5,"price_gbp":750,"label":"5 questionnaires"},
        "ten":{"qty":10,"price_gbp":1500,"label":"10 questionnaires"},
    }

try:
    openai_client = OpenAI(api_key=os.environ.get("OPENAI_API_KEY", "placeholder"))
except Exception as _oe:
    print(f"OpenAI client init warning: {_oe}")
    openai_client = None

try:
    claude_client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", "placeholder"))
except Exception as _ae:
    print(f"Anthropic client init warning: {_ae}")
    claude_client = None

# ── Constants (defined early so all routes can use them) ─────────────────────
DEMO_MAX_QUESTIONS = 10
DEMO_MAX_PER_TAB   = 3

SESSIONS: dict = {}
EMBEDDING_MODEL = "text-embedding-3-small"
CHUNK_SIZE = 400        # Smaller chunks = more precise retrieval
CHUNK_OVERLAP = 100     # More overlap = fewer gaps
TOP_K = 20              # Retrieve more chunks per question
EMBED_BATCH = 20
MAX_CHUNKS = 1200       # Allow more chunks = more document coverage
DOC_EXTENSIONS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls', '.csv', '.txt'}

app = FastAPI(
    title="CybeSure SecureAnswer",
    description="© CybeSure Ltd. All rights reserved. SecureAnswer™ is a trademark of CybeSure Ltd.",
    version="1.0.0"
)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

@app.middleware("http")
async def add_security_headers(request, call_next):
    response = await call_next(request)
    response.headers["X-Content-Type-Options"] = "nosniff"
    response.headers["X-Frame-Options"] = "DENY"
    response.headers["X-XSS-Protection"] = "1; mode=block"
    response.headers["Referrer-Policy"] = "strict-origin-when-cross-origin"
    response.headers["X-Powered-By"] = "CybeSure SecureAnswer"
    response.headers["X-Copyright"] = "© CybeSure Ltd. All rights reserved."
    response.headers["X-Data-Residency"] = "UK"
    response.headers["X-GDPR-Compliant"] = "true"
    return response

@app.on_event("startup")
async def startup():
    print(f"CybeSure SecureAnswer starting — AUTH_ENABLED={AUTH_ENABLED}")
    if AUTH_ENABLED:
        try:
            db = _SessionLocal()
            # Auto-create default admin if no users exist
            if db.query(User).count() == 0:
                seed_email = os.environ.get("ADMIN_EMAIL", "martfiddler@gmail.com")
                seed_pass  = os.environ.get("ADMIN_PASSWORD", "CybeSure2025!")
                seed_org   = os.environ.get("ADMIN_ORG", "CybeSure")
                slug = seed_org.lower().replace(" ", "-")
                org = Organisation(
                    name=seed_org, slug=slug,
                    contact_email=seed_email,
                    tier="test", status="active",
                    questionnaire_limit=999,
                    subscription_end=datetime.utcnow() + timedelta(days=3650)
                )
                db.add(org); db.flush()
                admin = User(
                    org_id=org.id, email=seed_email,
                    full_name="Martin Fiddler",
                    hashed_password=hash_password(seed_pass),
                    role="admin", is_active=True
                )
                db.add(admin); db.commit()
                print(f"✓ Auto-created admin account: {seed_email}")
            else:
                print(f"✓ {db.query(User).count()} users in database")
            db.close()
        except Exception as e:
            print(f"Startup seed error: {e}")

# All auth routes are defined inline below — no separate router needed

# ── Inline Auth Routes ────────────────────────────────────────────────────────

class RegisterRequest(BaseModel):
    org_name: str
    contact_email: str
    admin_name: str
    admin_email: str
    password: str

class InviteUserRequest(BaseModel):
    email: str
    full_name: str
    role: str = "contributor"
    password: str

@app.get("/debug/status")
async def debug_status(db=Depends(get_db)):
    """Shows DB status — remove after debugging."""
    status = {
        "auth_enabled": AUTH_ENABLED,
        "admin_email_env": os.environ.get("ADMIN_EMAIL", "NOT SET"),
        "admin_password_set": bool(os.environ.get("ADMIN_PASSWORD")),
    }
    if AUTH_ENABLED and db:
        try:
            status["user_count"] = db.query(User).count()
            status["org_count"] = db.query(Organisation).count()
            status["users"] = [
                {"email": u.email, "role": u.role, "active": u.is_active}
                for u in db.query(User).all()
            ]
        except Exception as e:
            status["db_error"] = str(e)
    return status


@app.get("/test-login", response_class=HTMLResponse)
async def test_login_page():
    """Simple test login page — bypasses the main frontend."""
    return HTMLResponse("""
    <html><body style="font-family:Arial;background:#0b1829;color:#fff;display:flex;align-items:center;justify-content:center;height:100vh;margin:0">
    <div style="background:#132030;padding:40px;border:1px solid #00c8e0;width:360px">
      <h2 style="color:#00c8e0;margin:0 0 24px">CybeSure — Test Login</h2>
      <div id="result" style="margin-bottom:16px;padding:10px;display:none"></div>
      <input id="email" type="email" value="martfiddler@gmail.com"
        style="width:100%;padding:10px;background:#0b1829;border:1px solid #00c8e0;color:#fff;margin-bottom:10px;box-sizing:border-box"/><br>
      <input id="pass" type="password" value="CybeSure2025!"
        style="width:100%;padding:10px;background:#0b1829;border:1px solid #00c8e0;color:#fff;margin-bottom:16px;box-sizing:border-box"/><br>
      <button onclick="tryLogin()"
        style="width:100%;background:#00c8e0;border:none;color:#0b1829;padding:12px;font-size:16px;font-weight:700;cursor:pointer">
        Test Login
      </button>
      <script>
      async function tryLogin() {
        const email = document.getElementById('email').value;
        const pass = document.getElementById('pass').value;
        const result = document.getElementById('result');
        try {
          const r = await fetch('/auth/login-json', {
            method: 'POST',
            headers: {'Content-Type':'application/json'},
            body: JSON.stringify({email, password: pass})
          });
          const d = await r.json();
          result.style.display = 'block';
          if(r.ok) {
            result.style.background = '#0a2a1a';
            result.style.border = '1px solid #00d4a0';
            result.style.color = '#00d4a0';
            result.innerHTML = '✅ LOGIN SUCCESS!<br>Token: ' + d.access_token.substring(0,20) + '...<br>User: ' + d.user_name + '<br>Tier: ' + d.subscription_tier;
          } else {
            result.style.background = '#2a0a0a';
            result.style.border = '1px solid #e84855';
            result.style.color = '#e84855';
            result.innerHTML = '❌ FAILED: ' + JSON.stringify(d);
          }
        } catch(e) {
          result.style.display = 'block';
          result.innerHTML = '❌ ERROR: ' + e.message;
        }
      }
      </script>
    </div></body></html>
    """)
async def reset_admin_password_get(db=Depends(get_db)):
    """Visit this URL to reset admin password."""
    try:
        if not AUTH_ENABLED or db is None:
            return HTMLResponse("<h1 style='color:red'>Auth not available — check Render logs</h1>")
        seed_email = os.environ.get("ADMIN_EMAIL", "martfiddler@gmail.com")
        seed_pass  = os.environ.get("ADMIN_PASSWORD", "CybeSure2025!")
        user = db.query(User).filter(User.email == seed_email).first()
        if not user:
            existing_org = db.query(Organisation).first()
            if existing_org:
                org = existing_org
            else:
                org = Organisation(
                    name="CybeSure", slug="cybesure",
                    contact_email=seed_email, tier="test", status="active",
                    questionnaire_limit=999,
                    subscription_end=datetime.utcnow() + timedelta(days=3650)
                )
                db.add(org); db.flush()
            user = User(
                org_id=org.id, email=seed_email, full_name="Martin Fiddler",
                hashed_password=hash_password(seed_pass),
                role="admin", is_active=True
            )
            db.add(user); db.commit()
            action = "CREATED"
        else:
            user.hashed_password = hash_password(seed_pass)
            user.is_active = True
            db.commit()
            action = "RESET"
        return HTMLResponse(f"""
        <html><body style="font-family:Arial;background:#0b1829;color:#fff;display:flex;align-items:center;justify-content:center;height:100vh;margin:0;flex-direction:column;gap:16px">
          <div style="font-size:48px">✅</div>
          <h2 style="color:#00d4a0;margin:0">Password {action} Successfully</h2>
          <div style="background:#132030;padding:20px 32px;border:1px solid #00c8e0;text-align:center">
            <div style="color:#7a9cbf;font-size:12px;margin-bottom:8px">LOGIN CREDENTIALS</div>
            <div style="font-size:18px;margin-bottom:4px"><strong>Email:</strong> {seed_email}</div>
            <div style="font-size:18px"><strong>Password:</strong> {seed_pass}</div>
          </div>
          <a href="/" style="color:#00c8e0;font-size:14px">← Go to Login</a>
        </body></html>
        """)
    except Exception as e:
        return HTMLResponse(f"<h1 style='color:red'>Error: {e}</h1><pre>{e}</pre>")


@app.post("/auth/reset-admin-password")
async def reset_admin_password_post(db=Depends(get_db)):
    """POST version of password reset."""
    return await reset_admin_password_get(db)


@app.post("/auth/bootstrap")
async def bootstrap(db=Depends(get_db)):
    """Create first admin account if DB is empty."""
    return await reset_admin_password_get(db)


@app.post("/auth/register")
async def register(req: RegisterRequest, db=Depends(get_db)):
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth system not available. Please try again shortly.")
    if db.query(User).filter(User.email == req.admin_email).first():
        raise HTTPException(400, "Email already registered")
    slug = req.org_name.lower().replace(" ", "-")[:50]
    base = slug; i = 1
    while db.query(Organisation).filter(Organisation.slug == slug).first():
        slug = f"{base}-{i}"; i += 1
    org = Organisation(
        name=req.org_name, slug=slug, contact_email=req.contact_email,
        tier="starter", status="trial", questionnaire_limit=50,
        subscription_end=datetime.utcnow() + timedelta(days=14)
    )
    db.add(org); db.flush()
    admin = User(
        org_id=org.id, email=req.admin_email, full_name=req.admin_name,
        hashed_password=hash_password(req.password), role="admin", is_active=True
    )
    db.add(admin); db.commit()
    return {"message": "Organisation registered. 14-day trial activated.", "org_id": org.id}


@app.post("/auth/register-demo")
async def register_demo(req: RegisterRequest, db=Depends(get_db)):
    """Register a demo account — limited to 3 questionnaires, 10 questions each."""
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth system not available.")
    existing = db.query(User).filter(User.email == req.admin_email).first()
    if existing:
        raise HTTPException(400, "Email already registered. Use Sign In.")
    slug = f"demo-{req.org_name.lower().replace(' ','-')[:30]}"
    base = slug; i = 1
    while db.query(Organisation).filter(Organisation.slug == slug).first():
        slug = f"{base}-{i}"; i += 1
    org = Organisation(
        name=req.org_name, slug=slug, contact_email=req.contact_email,
        tier="demo", status="active", questionnaire_limit=3,
        subscription_end=datetime.utcnow() + timedelta(days=30)
    )
    db.add(org); db.flush()
    admin = User(
        org_id=org.id, email=req.admin_email, full_name=req.admin_name,
        hashed_password=hash_password(req.password), role="admin", is_active=True
    )
    db.add(admin); db.commit()
    return {
        "message": "Demo account created. Max 3 questionnaires, 10 questions each.",
        "org_id": org.id, "tier": "demo"
    }


@app.post("/auth/activate-test")
async def activate_test(db=Depends(get_db), token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
    """Activate test mode for the current user's organisation — unlimited access."""
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth not available")
    user = decode_token(token)
    org = db.query(Organisation).filter(Organisation.id == user.org_id).first()
    org.tier = "test"
    org.status = "active"
    org.questionnaire_limit = 999
    org.questionnaires_used = 0
    org.topup_credits = 0
    org.subscription_end = datetime.utcnow() + timedelta(days=365)
    db.commit()
    return {"message": "Test mode activated. Unlimited access enabled.", "tier": "test"}

class LoginRequest(BaseModel):
    email: str
    password: str

@app.post("/auth/login-json")
async def login_json(req: LoginRequest, db=Depends(get_db)):
    """JSON login alternative — same as /auth/login but accepts JSON body."""
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth system not available.")
    user = db.query(User).filter(User.email == req.email).first()
    if not user or not verify_password(req.password, user.hashed_password):
        raise HTTPException(401, "Incorrect email or password")
    if not user.is_active:
        raise HTTPException(403, "Account inactive")
    org = db.query(Organisation).filter(Organisation.id == user.org_id).first()
    user.last_login = datetime.utcnow(); db.commit()
    token = make_token({"sub": str(user.id), "org_id": user.org_id})
    cfg = SUBSCRIPTION_CONFIG.get(org.tier, {})
    is_demo = org.tier == "demo"
    is_test = org.tier == "test"
    return {
        "access_token": token, "token_type": "bearer",
        "user_id": user.id, "user_name": user.full_name,
        "user_role": user.role, "org_id": org.id, "org_name": org.name,
        "subscription_tier": org.tier,
        "tier_name": cfg.get("name", org.tier),
        "questionnaires_remaining": org.remaining,
        "questionnaires_used": org.questionnaires_used,
        "questionnaire_limit": org.total_limit,
        "subscription_expires": org.subscription_end.isoformat(),
        "is_demo": is_demo, "is_test": is_test,
        "demo_q_limit": DEMO_MAX_QUESTIONS if is_demo else None,
        "woocommerce_verified": bool(WC_CONSUMER_KEY)
    }


@app.post("/auth/login")
async def login(form: OAuth2PasswordRequestForm = Depends(), db=Depends(get_db)):
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth system not available. Please try again shortly.")
    user = db.query(User).filter(User.email == form.username).first()
    if not user or not verify_password(form.password, user.hashed_password):
        raise HTTPException(401, "Incorrect email or password")
    if not user.is_active:
        raise HTTPException(403, "Account inactive")
    org = db.query(Organisation).filter(Organisation.id == user.org_id).first()

    # Check WooCommerce subscription — upgrade tier if paid subscription found
    if WC_CONSUMER_KEY:
        wc_tier = get_woocommerce_subscription(user.email)
        if wc_tier and wc_tier != org.tier:
            cfg = SUBSCRIPTION_CONFIG.get(wc_tier, {})
            org.tier = wc_tier
            org.status = "active"
            org.questionnaire_limit = cfg.get("limit", org.questionnaire_limit)
            if org.subscription_end < datetime.utcnow():
                org.subscription_end = datetime.utcnow() + timedelta(days=365)
            db.commit()
            print(f"WooCommerce upgrade: {user.email} → {wc_tier}")

    user.last_login = datetime.utcnow(); db.commit()
    token = make_token({"sub": str(user.id), "org_id": user.org_id})
    cfg = SUBSCRIPTION_CONFIG.get(org.tier, {})
    is_demo = org.tier == "demo"
    is_test = org.tier == "test"
    return {
        "access_token": token, "token_type": "bearer",
        "user_id": user.id, "user_name": user.full_name,
        "user_role": user.role, "org_id": org.id, "org_name": org.name,
        "subscription_tier": org.tier,
        "tier_name": cfg.get("name", org.tier),
        "questionnaires_remaining": org.remaining,
        "questionnaires_used": org.questionnaires_used,
        "questionnaire_limit": org.total_limit,
        "subscription_expires": org.subscription_end.isoformat(),
        "is_demo": is_demo,
        "is_test": is_test,
        "demo_q_limit": DEMO_MAX_QUESTIONS if is_demo else None,
        "woocommerce_verified": bool(WC_CONSUMER_KEY)
    }

@app.get("/auth/me")
async def get_me(db=Depends(get_db), token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth not available")
    user = decode_token(token)
    org = db.query(Organisation).filter(Organisation.id == user.org_id).first()
    return {
        "user_id": user.id, "email": user.email, "full_name": user.full_name,
        "role": user.role, "org_id": org.id, "org_name": org.name,
        "tier": org.tier, "status": org.status,
        "remaining": org.remaining, "total_limit": org.total_limit,
        "used": org.questionnaires_used,
        "subscription_end": org.subscription_end.isoformat(),
        "is_active": org.is_active
    }

@app.get("/org/dashboard")
async def org_dashboard(db=Depends(get_db), token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth not available")
    user = decode_token(token)
    org = db.query(Organisation).filter(Organisation.id == user.org_id).first()
    users_count = db.query(User).filter(User.org_id==org.id, User.is_active==True).count()
    recent = db.query(QuestionnaireRun).filter(
        QuestionnaireRun.org_id==org.id
    ).order_by(QuestionnaireRun.created_at.desc()).limit(10).all()
    cfg = SUBSCRIPTION_CONFIG.get(org.tier, {})
    days_left = max(0, (org.subscription_end - datetime.utcnow()).days)
    return {
        "org_name": org.name, "tier": org.tier, "tier_name": cfg.get("name", org.tier),
        "status": org.status, "is_active": org.is_active,
        "questionnaire_limit": org.questionnaire_limit,
        "questionnaires_used": org.questionnaires_used,
        "topup_credits": org.topup_credits, "remaining": org.remaining,
        "total_limit": org.total_limit,
        "subscription_end": org.subscription_end.isoformat(),
        "days_remaining": days_left, "users_count": users_count,
        "recent_runs": [{"id":r.id,"filename":r.filename,"status":r.status,
                        "created_at":r.created_at.isoformat()} for r in recent],
        "subscription_plans": SUBSCRIPTION_CONFIG,
        "topup_options": {"single":{"qty":1,"price_gbp":300},"bundle":{"qty":10,"price_gbp":1000}}
    }

@app.post("/admin/users/invite")
async def invite_user(req: InviteUserRequest, db=Depends(get_db), token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth not available")
    user = decode_token(token)
    if user.role != "admin":
        raise HTTPException(403, "Admin access required")
    if db.query(User).filter(User.email == req.email).first():
        raise HTTPException(400, "Email already registered")
    new_user = User(
        org_id=user.org_id, email=req.email, full_name=req.full_name,
        hashed_password=hash_password(req.password), role=req.role, is_active=True
    )
    db.add(new_user); db.commit(); db.refresh(new_user)
    return {"message": f"User {req.full_name} added as {req.role}", "user_id": new_user.id}

@app.get("/admin/users")
async def list_users(db=Depends(get_db), token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth not available")
    user = decode_token(token)
    if user.role != "admin":
        raise HTTPException(403, "Admin access required")
    users = db.query(User).filter(User.org_id == user.org_id).all()
    return [{"id":u.id,"email":u.email,"full_name":u.full_name,"role":u.role,
             "is_active":u.is_active,"created_at":u.created_at.isoformat(),
             "last_login":u.last_login.isoformat() if u.last_login else None} for u in users]

@app.delete("/admin/users/{user_id}")
async def delete_user(user_id: int, db=Depends(get_db), token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth not available")
    admin = get_auth_user(token)
    if admin.role != "admin":
        raise HTTPException(403, "Admin access required")
    user = db.query(User).filter(User.id == user_id, User.org_id == admin.org_id).first()
    if not user:
        raise HTTPException(404, "User not found")
    if user.id == admin.id:
        raise HTTPException(400, "Cannot remove yourself")
    user.is_active = False
    db.commit()
    return {"message": "User deactivated"}


@app.get("/analytics/dashboard")
async def analytics_dashboard(db=Depends(get_db), token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
    """Full analytics dashboard with time savings calculations."""
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth not available")
    user = decode_token(token)
    org_id = user.org_id

    runs = db.query(QuestionnaireRun).filter(
        QuestionnaireRun.org_id == org_id
    ).order_by(QuestionnaireRun.created_at.desc()).all()

    # Time savings constants (minutes per question)
    MINS_SIMPLE  = 5
    MINS_AVERAGE = 10
    MINS_COMPLEX = 15

    total_runs = len(runs)
    completed_runs = [r for r in runs if r.status == "complete"]
    total_questions = sum(r.question_count or 0 for r in runs)
    avg_questions = round(total_questions / total_runs, 1) if total_runs else 0

    # Processing times
    proc_times = [r.processing_secs for r in completed_runs if r.processing_secs and r.processing_secs > 0]
    avg_proc_secs = round(sum(proc_times) / len(proc_times), 1) if proc_times else 0
    avg_proc_per_q = round(avg_proc_secs / avg_questions, 1) if avg_questions and avg_proc_secs else 0

    # Time savings
    manual_mins_simple  = total_questions * MINS_SIMPLE
    manual_mins_average = total_questions * MINS_AVERAGE
    manual_mins_complex = total_questions * MINS_COMPLEX
    ai_mins_total       = round(sum(proc_times) / 60, 1)

    saved_simple  = manual_mins_simple  - ai_mins_total
    saved_average = manual_mins_average - ai_mins_total
    saved_complex = manual_mins_complex - ai_mins_total

    # Industry breakdown
    sector_counts = {}
    for r in runs:
        sector = r.industry_sector or "Not specified"
        sector_counts[sector] = sector_counts.get(sector, 0) + 1
    sectors = sorted(sector_counts.items(), key=lambda x: x[1], reverse=True)

    # Monthly trend (last 12 months)
    from collections import defaultdict
    monthly = defaultdict(int)
    for r in runs:
        key = r.created_at.strftime("%b %Y")
        monthly[key] += 1

    # Recent runs
    recent = []
    for r in runs[:10]:
        recent.append({
            "id": r.id,
            "filename": r.filename or "Portal questionnaire",
            "question_count": r.question_count or 0,
            "status": r.status,
            "industry_sector": r.industry_sector or "Not specified",
            "processing_secs": r.processing_secs or 0,
            "created_at": r.created_at.isoformat(),
            "completed_at": r.completed_at.isoformat() if r.completed_at else None
        })

    return {
        "summary": {
            "total_questionnaires": total_runs,
            "completed_questionnaires": len(completed_runs),
            "total_questions_answered": total_questions,
            "avg_questions_per_questionnaire": avg_questions,
            "avg_processing_secs_per_questionnaire": avg_proc_secs,
            "avg_processing_secs_per_question": avg_proc_per_q,
        },
        "time_savings": {
            "ai_processing_mins": ai_mins_total,
            "manual_mins_simple_scenario":  round(manual_mins_simple, 1),
            "manual_mins_average_scenario": round(manual_mins_average, 1),
            "manual_mins_complex_scenario": round(manual_mins_complex, 1),
            "saved_mins_simple":  round(saved_simple, 1),
            "saved_mins_average": round(saved_average, 1),
            "saved_mins_complex": round(saved_complex, 1),
            "saved_hrs_simple":   round(saved_simple / 60, 1),
            "saved_hrs_average":  round(saved_average / 60, 1),
            "saved_hrs_complex":  round(saved_complex / 60, 1),
            "per_question_saved_average": MINS_AVERAGE,
            "per_questionnaire_saved_average": round(avg_questions * MINS_AVERAGE, 1),
        },
        "industry_sectors": [{"sector": s, "count": c} for s, c in sectors],
        "monthly_trend": [{"month": k, "count": v} for k, v in list(monthly.items())[-12:]],
        "recent_runs": recent
    }

@app.post("/billing/create-topup-session")
async def create_topup(topup_type: str, db=Depends(get_db), token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
    """Create Stripe checkout for questionnaire top-ups."""
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth not available")
    user = decode_token(token)
    options = {
        "five":   {"qty": 5,  "price_gbp": 750,  "label": "5 additional questionnaires"},
        "ten":    {"qty": 10, "price_gbp": 1500, "label": "10 additional questionnaires"},
        "single": {"qty": 1,  "price_gbp": 300,  "label": "1 additional questionnaire"},
        "bundle": {"qty": 10, "price_gbp": 1000, "label": "10 questionnaire bundle"},
    }
    opt = options.get(topup_type)
    if not opt: raise HTTPException(400, "Invalid topup type. Use: five, ten")
    stripe_key = os.environ.get("STRIPE_SECRET_KEY", "")
    if not stripe_key:
        raise HTTPException(503, "Payment not configured. Contact support@cybesure.com")
    try:
        import stripe
        stripe.api_key = stripe_key
        org = db.query(Organisation).filter(Organisation.id == user.org_id).first()
        app_url = os.environ.get("APP_URL", "https://cybesure-secureanswer-2.onrender.com")
        session = stripe.checkout.Session.create(
            payment_method_types=["card"],
            line_items=[{"price_data": {"currency": "gbp",
                "product_data": {"name": f"CybeSure SecureAnswer — {opt['label']}"},
                "unit_amount": int(opt['price_gbp'] * 100)}, "quantity": 1}],
            mode="payment",
            success_url=f"{app_url}/billing/success?session_id={{CHECKOUT_SESSION_ID}}",
            cancel_url=f"{app_url}/billing/cancel",
            customer_email=user.email,
            metadata={"org_id": str(org.id), "topup_type": topup_type,
                      "qty": str(opt['qty']), "user_id": str(user.id)}
        )
        return {"checkout_url": session.url}
    except Exception as e:
        raise HTTPException(400, str(e))


@app.post("/billing/activate-test-mode")
async def activate_test_mode(db=Depends(get_db), token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
    """Activate test mode — gives unlimited access for testing all features."""
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth not available")
    user = decode_token(token)
    org = db.query(Organisation).filter(Organisation.id == user.org_id).first()
    prev_tier = org.tier
    org.tier = "test"
    org.status = "active"
    org.questionnaire_limit = 999
    org.questionnaires_used = 0
    org.topup_credits = 0
    org.subscription_end = datetime.utcnow() + timedelta(days=365)
    db.commit()
    return {
        "message": f"Test mode activated (was: {prev_tier}). Unlimited questionnaires enabled.",
        "tier": "test",
        "limit": 999
    }


@app.post("/billing/activate-demo")
async def activate_demo(db=Depends(get_db), token: str = Depends(OAuth2PasswordBearer(tokenUrl="/auth/login"))):
    """Switch to demo mode to test the demo experience."""
    if not AUTH_ENABLED or db is None:
        raise HTTPException(503, "Auth not available")
    user = decode_token(token)
    org = db.query(Organisation).filter(Organisation.id == user.org_id).first()
    org.tier = "demo"
    org.status = "active"
    org.questionnaire_limit = 3
    org.questionnaires_used = 0
    org.topup_credits = 0
    org.subscription_end = datetime.utcnow() + timedelta(days=30)
    db.commit()
    return {"message": "Demo mode activated. Max 3 questionnaires, 10 questions each.", "tier": "demo"}


@app.get("/billing/success", response_class=HTMLResponse)
async def billing_success():
    return HTMLResponse("<html><body style='font-family:Arial;background:#0b1829;color:#fff;display:flex;align-items:center;justify-content:center;height:100vh;margin:0'><div style='text-align:center'><div style='font-size:48px'>✅</div><h2 style='color:#00d4a0'>Payment Successful</h2><p style='color:#7a9cbf'>Credits added to your account.</p><a href='/' style='color:#00c8e0'>Return to SecureAnswer</a></div></body></html>")

@app.get("/billing/cancel", response_class=HTMLResponse)
async def billing_cancel():
    return HTMLResponse("<html><body style='font-family:Arial;background:#0b1829;color:#fff;display:flex;align-items:center;justify-content:center;height:100vh;margin:0'><div style='text-align:center'><div style='font-size:48px'>❌</div><h2 style='color:#f5a623'>Payment Cancelled</h2><a href='/' style='color:#00c8e0'>Return to SecureAnswer</a></div></body></html>")

@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    index_path = os.path.join(os.path.dirname(__file__), "index.html")
    if os.path.exists(index_path):
        with open(index_path, "r") as f:
            return HTMLResponse(content=f.read())
    return HTMLResponse(content="<h1>CybeSure SecureAnswer is running</h1>")

# ── questionnaire extraction ──────────────────────────────────────────────────

# Store questionnaire structure for matched export
QUESTIONNAIRE_STRUCTURE: dict = {}

def find_question_col(cols: list) -> str:
    """Find question column by header name."""
    cols_l = [str(c).strip().lower() for c in cols]
    for kw in ('question','questions','requirement','requirements','control','controls','query','ask'):
        if kw in cols_l:
            return cols[cols_l.index(kw)]
    for col, col_l in zip(cols, cols_l):
        if any(kw in col_l for kw in ('question','requirement','control')):
            return col
    return None

def find_answer_col(cols: list) -> str:
    """Find answer column by header name."""
    cols_l = [str(c).strip().lower() for c in cols]
    for kw in ('answer','answers','response','responses','vendor response','supplier response','your response','your answer'):
        if kw in cols_l:
            return cols[cols_l.index(kw)]
    for col, col_l in zip(cols, cols_l):
        if any(kw in col_l for kw in ('answer','response')):
            return col
    return None

def find_question_col_by_content(df) -> str:
    """Find question column by content analysis — longest text with most question marks."""
    best_col, best_score = None, 0
    for col in df.columns:
        vals = df[col].dropna().astype(str)
        score = sum(1 for v in vals if len(v) > 15) + sum(2 for v in vals if '?' in v)
        if score > best_score:
            best_score = score
            best_col = col
    return best_col


def extract_questions_from_excel(data: bytes) -> list:
    """
    Smart Excel extractor — reads ALL sheets, finds question column
    by name or content, handles any multi-column structure.
    """
    all_questions = []
    for engine in ['openpyxl', 'xlrd', None]:
        try:
            kwargs = {'engine': engine} if engine else {}
            xl = pd.ExcelFile(io.BytesIO(data), **kwargs)
            for sheet_name in xl.sheet_names:
                try:
                    rk = {'sheet_name': sheet_name}
                    if engine: rk['engine'] = engine
                    df = pd.read_excel(io.BytesIO(data), **rk)
                    if df.empty: continue
                    df.columns = [str(c).strip() for c in df.columns]
                    q_col = find_question_col(list(df.columns)) or find_question_col_by_content(df)
                    if q_col is None: continue
                    sheet_qs = []
                    for val in df[q_col].dropna():
                        text = str(val).strip()
                        if text and len(text) > 3 and text.lower() not in ['nan','none','n/a','-','question','questions','requirement','control']:
                            sheet_qs.append(text)
                    if sheet_qs:
                        all_questions.extend(sheet_qs)
                        print(f"Sheet '{sheet_name}': {len(sheet_qs)} questions from col '{q_col}'")
                except Exception as e:
                    print(f"Sheet '{sheet_name}' error: {e}")
            if all_questions: break
        except Exception as e:
            print(f"Engine {engine} failed: {e}")
    seen, unique = set(), []
    for q in all_questions:
        if q not in seen:
            seen.add(q)
            unique.append(q)
    return unique


def extract_questionnaire_structure(data: bytes, filename: str) -> Dict:
    """
    Extract full structure of questionnaire including all columns and all sheets.
    Returns structure dict used for matched export.
    """
    structure = {"filename": filename, "sheets": []}
    for engine in ['openpyxl', 'xlrd', None]:
        try:
            kwargs = {'engine': engine} if engine else {}
            xl = pd.ExcelFile(io.BytesIO(data), **kwargs)
            for sheet_name in xl.sheet_names:
                try:
                    rk = {'sheet_name': sheet_name}
                    if engine: rk['engine'] = engine
                    df = pd.read_excel(io.BytesIO(data), **rk)
                    if df.empty: continue
                    df.columns = [str(c).strip() for c in df.columns]
                    cols = list(df.columns)
                    q_col = find_question_col(cols) or find_question_col_by_content(df)
                    a_col = find_answer_col(cols)
                    sheet_info = {"name": sheet_name, "columns": cols, "question_col": q_col, "answer_col": a_col, "rows": []}
                    for idx, row in df.iterrows():
                        q_val = str(row.get(q_col, "")).strip() if q_col else ""
                        if q_val and len(q_val) > 3 and q_val.lower() not in ['nan','none','n/a']:
                            row_data = {col: (str(row[col]) if pd.notna(row[col]) else "") for col in cols}
                            sheet_info["rows"].append({"row_idx": int(idx), "question": q_val, "row_data": row_data})
                    structure["sheets"].append(sheet_info)
                except Exception as e:
                    print(f"Structure extract error sheet '{sheet_name}': {e}")
            if structure["sheets"]: break
        except Exception as e:
            print(f"Structure engine {engine} failed: {e}")
    return structure

def extract_questions_from_csv(data: bytes) -> List[str]:
    """Extract questions from CSV files."""
    questions = []
    try:
        for enc in ['utf-8', 'latin-1', 'cp1252']:
            try:
                df = pd.read_csv(io.BytesIO(data), encoding=enc)
                break
            except Exception:
                continue

        df.columns = [str(c).strip().lower() for c in df.columns]

        q_col = None
        for col in df.columns:
            if col in ('question', 'questions', 'requirement', 'requirements', 'control', 'description'):
                q_col = col
                break

        if q_col is None:
            best = 0
            for col in df.columns:
                vals = df[col].dropna().astype(str)
                score = sum(1 for v in vals if len(v) > 15)
                if score > best:
                    best = score
                    q_col = col

        if q_col:
            for val in df[q_col].dropna():
                text = str(val).strip()
                if text and len(text) > 3 and text.lower() not in ['nan', 'none', 'n/a']:
                    questions.append(text)

    except Exception as e:
        print(f"CSV parse error: {e}")
    return questions


def extract_questions_from_docx(data: bytes) -> List[str]:
    """Extract questions from Word documents."""
    questions = []
    try:
        doc = Document(io.BytesIO(data))
        for p in doc.paragraphs:
            text = p.text.strip()
            if text and len(text) > 3:
                questions.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and len(text) > 3:
                        questions.append(text)
    except Exception as e:
        print(f"DOCX parse error: {e}")
    return questions


def extract_questions_from_pdf(data: bytes) -> List[str]:
    """Extract questions from PDF files."""
    questions = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    for line in text.split('\n'):
                        line = line.strip()
                        if line and len(line) > 5:
                            questions.append(line)
    except Exception as e:
        print(f"PDF parse error: {e}")
    return questions


def extract_questions(filename: str, data: bytes) -> List[str]:
    """Route to correct extractor."""
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext in ("xlsx", "xls"):
        return extract_questions_from_excel(data)
    elif ext == "csv":
        return extract_questions_from_csv(data)
    elif ext in ("docx", "doc"):
        return extract_questions_from_docx(data)
    elif ext == "pdf":
        return extract_questions_from_pdf(data)
    else:
        return [l.strip() for l in data.decode("utf-8", errors="ignore").split('\n') if l.strip() and len(l.strip()) > 5]

# ── document file parsing ─────────────────────────────────────────────────────

def parse_pdf(data: bytes) -> List[str]:
    pages = []
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            for p in pdf.pages:
                t = p.extract_text()
                if t and t.strip():
                    pages.append(t[:4000])
    except Exception as e:
        print(f"PDF error: {e}")
    return pages

def parse_docx(data: bytes) -> List[str]:
    try:
        doc = Document(io.BytesIO(data))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    texts.append(row_text)
        return texts
    except Exception as e:
        print(f"DOCX error: {e}")
        return []

def parse_doc(data: bytes) -> List[str]:
    try:
        doc = Document(io.BytesIO(data))
        texts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    texts.append(row_text)
        if texts:
            return texts
    except Exception:
        pass
    try:
        text = data.decode('latin-1', errors='ignore')
        chunks = re.findall(r'[A-Za-z0-9\s\.\,\!\?\:\;\-\(\)\'\"]{20,}', text)
        readable = ' '.join(chunks)
        if readable.strip():
            return [readable[:6000]]
    except Exception:
        pass
    return []

def parse_excel_doc(data: bytes) -> List[str]:
    try:
        # Read all sheets
        xl = pd.ExcelFile(io.BytesIO(data))
        all_rows = []
        for sheet in xl.sheet_names:
            df = pd.read_excel(io.BytesIO(data), sheet_name=sheet)
            for _, row in df.iterrows():
                r = " | ".join(str(v) for v in row.values if pd.notna(v) and str(v).strip())
                if r.strip():
                    all_rows.append(r)
        return all_rows
    except Exception as e:
        print(f"Excel error: {e}")
        return []

def parse_csv_doc(data: bytes) -> List[str]:
    try:
        df = pd.read_csv(io.BytesIO(data))
        rows = []
        for _, row in df.iterrows():
            r = " | ".join(str(v) for v in row.values if pd.notna(v) and str(v).strip())
            if r.strip():
                rows.append(r)
        return rows
    except Exception as e:
        print(f"CSV error: {e}")
        return []

def parse_file(filename: str, data: bytes) -> List[str]:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf": return parse_pdf(data)
    elif ext == "docx": return parse_docx(data)
    elif ext == "doc": return parse_doc(data)
    elif ext in ("xlsx", "xls"): return parse_excel_doc(data)
    elif ext == "csv": return parse_csv_doc(data)
    else: return [data.decode("utf-8", errors="ignore")[:5000]]

# ── URL fetching ──────────────────────────────────────────────────────────────

def is_document_url(url: str) -> bool:
    path = urlparse(url).path.lower()
    return any(path.endswith(ext) for ext in DOC_EXTENSIONS)

def fetch_document_from_url(url: str) -> Tuple[str, bytes]:
    headers = {'User-Agent': 'Mozilla/5.0 (compatible; CybeSure/1.0)', 'Accept': '*/*'}
    resp = requests.get(url, headers=headers, timeout=30, allow_redirects=True)
    resp.raise_for_status()
    filename = url.split('/')[-1].split('?')[0] or 'document'
    cd = resp.headers.get('Content-Disposition', '')
    if 'filename=' in cd:
        filename = cd.split('filename=')[-1].strip().strip('"\'')
    if '.' not in filename:
        ct = resp.headers.get('Content-Type', '')
        if 'pdf' in ct: filename += '.pdf'
        elif 'word' in ct: filename += '.docx'
        elif 'excel' in ct: filename += '.xlsx'
    return filename, resp.content

def discover_documents_from_page(url: str) -> List[str]:
    headers = {'User-Agent': 'Mozilla/5.0 (compatible; CybeSure/1.0)'}
    try:
        resp = requests.get(url, headers=headers, timeout=20, allow_redirects=True)
        resp.raise_for_status()
    except Exception as e:
        print(f"Failed to fetch page {url}: {e}")
        return []
    soup = BeautifulSoup(resp.text, 'html.parser')
    base = f"{urlparse(url).scheme}://{urlparse(url).netloc}"
    doc_urls = []
    for tag in soup.find_all('a', href=True):
        href = tag['href']
        if href.startswith('http'): full_url = href
        elif href.startswith('/'): full_url = base + href
        else: full_url = urljoin(url, href)
        if is_document_url(full_url):
            doc_urls.append(full_url)
    return list(set(doc_urls))

def fetch_all_from_url(url: str) -> List[Tuple[str, bytes]]:
    results = []
    # Google Drive single file
    if 'drive.google.com/file' in url or ('docs.google.com' in url and '/d/' in url):
        file_id = None
        if '/d/' in url: file_id = url.split('/d/')[1].split('/')[0]
        elif 'id=' in url: file_id = url.split('id=')[1].split('&')[0]
        if file_id:
            try:
                fname, data = fetch_document_from_url(
                    f"https://drive.google.com/uc?export=download&id={file_id}")
                results.append((fname, data))
                return results
            except Exception as e:
                print(f"Google Drive download failed: {e}")
    # Direct document URL
    if is_document_url(url):
        try:
            fname, data = fetch_document_from_url(url)
            results.append((fname, data))
        except Exception as e:
            print(f"Failed to fetch {url}: {e}")
        return results
    # Page scan
    try:
        doc_urls = discover_documents_from_page(url)
        if doc_urls:
            for doc_url in doc_urls[:20]:
                try:
                    fname, data = fetch_document_from_url(doc_url)
                    results.append((fname, data))
                except Exception as e:
                    print(f"Failed {doc_url}: {e}")
        else:
            resp = requests.get(url, timeout=20)
            soup = BeautifulSoup(resp.text, 'html.parser')
            text = soup.get_text(separator='\n', strip=True)
            if text:
                results.append(('webpage.txt', text.encode('utf-8')))
    except Exception as e:
        print(f"Page scan failed: {e}")
    return results

# ── chunking ──────────────────────────────────────────────────────────────────

def simple_chunk(text: str) -> List[str]:
    words = text.split()
    chunks, start = [], 0
    while start < len(words):
        end = min(start + CHUNK_SIZE, len(words))
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk)
        if end == len(words):
            break
        start += CHUNK_SIZE - CHUNK_OVERLAP
    return chunks

# ── embeddings ────────────────────────────────────────────────────────────────

def embed_texts(texts: List[str]) -> List[List[float]]:
    if openai_client is None:
        raise HTTPException(503, "OpenAI API key not configured. Add OPENAI_API_KEY to environment variables.")
    all_emb = []
    for i in range(0, len(texts), EMBED_BATCH):
        resp = openai_client.embeddings.create(
            input=texts[i:i+EMBED_BATCH], model=EMBEDDING_MODEL)
        all_emb.extend([r.embedding for r in resp.data])
        gc.collect()
    return all_emb

# ── retrieval ─────────────────────────────────────────────────────────────────

def expand_query(q: str) -> str:
    """Expand query with cyber security domain terms for better retrieval."""
    return (f"{q} policy procedure control implementation process "
            f"cyber security information security governance compliance "
            f"management system requirement standard evidence")

def retrieve(session_id: str, question: str) -> List[str]:
    """Multi-angle retrieval for maximum document coverage."""
    index, chunks, _ = SESSIONS[session_id]  # unpack 3-tuple

    queries = [
        expand_query(question),
        question,
        question.lower().replace('?','').replace('do you','').replace('how do you','').strip(),
    ]

    seen = set()
    results = []
    for query in queries:
        try:
            qv = np.array([embed_texts([query])[0]], dtype="float32")
            k = min(TOP_K, len(chunks))
            distances, idxs = index.search(qv, k)
            for dist, i in zip(distances[0], idxs[0]):
                c = chunks[i]
                if c not in seen:
                    seen.add(c)
                    results.append((dist, c))
        except Exception as e:
            print(f"Retrieval error: {e}")
            continue

    results.sort(key=lambda x: x[0])
    return [c for _, c in results[:TOP_K]]

# ── Claude ────────────────────────────────────────────────────────────────────

SYSTEM_PROMPT = """You are a senior cyber security compliance expert writing comprehensive answers to a security questionnaire on behalf of an organisation.

You will be given extracts from the organisation's security policies and documents. Synthesise ALL relevant information across ALL provided extracts to write thorough, confident, professional answers — exactly as the organisation would respond to an auditor or customer due diligence questionnaire.

CRITICAL RULES:
- Always write as "We" — you ARE the organisation
- NEVER mention document names, policy titles, file names or "the documents"
- NEVER say what is missing, not covered, or not evidenced
- NEVER say "based on available documentation" or any hedging phrase
- NEVER reference "the extracts" or "the materials"
- Draw on ALL relevant content across every provided extract
- Be comprehensive — cover every relevant aspect, process, control and detail found
- Write with confidence and authority like a seasoned CISO responding to a customer
- Include specific details: tools, frequencies, roles, processes, technical controls

ANSWER STRUCTURE — always follow this exact 5-part pattern:

1. CONTROL STATEMENT — One confident opening sentence stating what we do
2. IMPLEMENTATION DETAIL — Comprehensive how: specific tools, processes, frequencies, technical controls, procedures, thresholds
3. GOVERNANCE & OWNERSHIP — Who owns and oversees this (specific roles/functions/teams)
4. STANDARDS ALIGNMENT — Frameworks this aligns to (only include if evidenced in the material: ISO 27001, NIST, GDPR, PCI DSS, Cyber Essentials etc)
5. EVIDENCE — Specific records, logs, reports, registers or artefacts that demonstrate this control

GOLD STANDARD EXAMPLE:
Q: "How do you control user access to systems?"
A: "We operate a role-based access control (RBAC) model aligned to least privilege principles across all systems and environments. Access is provisioned via a formal joiner-mover-leaver (JML) process integrated with HR workflows, with all access requests requiring documented approval from system owners and enforced through centralised identity management. Privileged access is strictly controlled with multi-factor authentication enforced for all administrative accounts, and standing privileged access is minimised through just-in-time provisioning. Access reviews are conducted quarterly with system owners required to validate user access appropriateness, and any discrepancies are remediated within defined SLAs. Ownership of access control sits with the Head of IT Operations with oversight from the Information Security function. This aligns with ISO 27001:2022 Annex A.5.15 / A.8.2, NIST AC family controls and PCI DSS v4 Requirement 7. Evidence includes access control policy records, IAM system audit logs, JML workflow records, quarterly access review reports and privilege access management logs."

Respond in JSON only — no markdown fences:
{
  "confidence": "Yes|No|Partial",
  "confidence_pct": 85,
  "explanation": "Comprehensive structured answer following the 5-part pattern — written confidently as the organisation with maximum detail",
  "sources": ["Specific control, process, tool or detail from the source material that supports this answer"]
}"""


IMPROVED_PROMPT = """You are a senior cyber security compliance consultant reviewing a draft answer to a security questionnaire question.

Your task is to take the existing answer and improve it to score 95%+ in a formal security assessment.

To achieve 95%+, the improved answer must:
1. Be more specific about technical controls, tools, configurations and thresholds
2. Reference additional relevant security frameworks and standards
3. Strengthen governance language with named roles and oversight structures  
4. Add more specific evidence types that would satisfy an auditor
5. Use more precise, quantified language (e.g. specific timeframes, frequency, SLAs)
6. Ensure every sentence adds audit value

RULES:
- Write as "We" — first person organisational voice
- NEVER mention documents or policy names
- NEVER say what is missing
- Build on the existing answer — do not replace it entirely
- Make it more detailed, more specific, more evidenced

Respond in JSON only:
{
  "improved_answer": "The enhanced answer that would score 95%+ in a formal assessment",
  "improvement_notes": "2-3 bullet points explaining what was strengthened and why it scores higher"
}"""


def ask_claude(question: str, chunks: List[str]) -> Dict:
    """Generate comprehensive answer plus an improved 95%+ version."""
    context = "\n\n---\n\n".join(chunks)

    # Step 1: Generate primary answer
    msg = claude_client.messages.create(
        model="claude-opus-4-5",
        max_tokens=1500,
        system=SYSTEM_PROMPT,
        messages=[{
            "role": "user",
            "content": (
                f"Organisation's security policy extracts:\n\n{context}\n\n"
                f"---\n\n"
                f"Security questionnaire question: {question}\n\n"
                f"Write a comprehensive, detailed answer as the organisation. "
                f"Use ALL relevant information from every extract above. "
                f"Follow the 5-part structure precisely. "
                f"Be specific about processes, tools, frequencies and controls. "
                f"Respond in JSON only."
            )
        }]
    )
    text = msg.content[0].text.strip().replace("```json", "").replace("```", "").strip()
    result = json.loads(text)
    if "confidence_pct" not in result:
        result["confidence_pct"] = {"Yes": 90, "Partial": 65, "No": 15}.get(
            result.get("confidence", "Partial"), 50)

    # Step 2: Generate improved 95%+ version
    try:
        improved_msg = claude_client.messages.create(
            model="claude-opus-4-5",
            max_tokens=1200,
            system=IMPROVED_PROMPT,
            messages=[{
                "role": "user",
                "content": (
                    f"Question: {question}\n\n"
                    f"Current answer:\n{result.get('explanation', '')}\n\n"
                    f"Additional context from organisation's documents:\n{context[:3000]}\n\n"
                    f"Improve this answer to score 95%+. Respond in JSON only."
                )
            }]
        )
        imp_text = improved_msg.content[0].text.strip().replace("```json","").replace("```","").strip()
        imp_data = json.loads(imp_text)
        result["improved_answer"] = str(imp_data.get("improved_answer", ""))
        # improvement_notes may be a list or string — always convert to string
        notes = imp_data.get("improvement_notes", "")
        if isinstance(notes, list):
            notes = "\n".join(str(n) for n in notes)
        result["improvement_notes"] = str(notes)
    except Exception as e:
        print(f"Improved answer error: {e}")
        result["improved_answer"] = ""
        result["improvement_notes"] = ""

    return result

# ── index builder ─────────────────────────────────────────────────────────────

# Answer cache: session_id -> {question_text: answer_dict}
ANSWER_CACHE: dict = {}

# Persistent question cache: question_hash -> answer_dict
# Survives across sessions — same question never processed twice
QUESTION_CACHE: dict = {}

def question_hash(question: str) -> str:
    """Create a stable hash key for a question."""
    return hashlib.sha256(question.strip().lower().encode()).hexdigest()[:16]

# Approval store: token -> approval_record
APPROVALS: dict = {}

# Email config from environment
SMTP_HOST = os.environ.get("SMTP_HOST", "smtp.gmail.com")
SMTP_PORT = int(os.environ.get("SMTP_PORT", "587"))
SMTP_USER = os.environ.get("SMTP_USER", "")
SMTP_PASS = os.environ.get("SMTP_PASS", "")
APP_URL   = os.environ.get("APP_URL", "https://cybesure-qa-platform.onrender.com")

def build_index(raw_texts: List[str], session_id: str, doc_names: List[str] = None) -> Dict:
    chunks = []
    for text in raw_texts:
        if text.strip():
            chunks.extend(simple_chunk(text))
    if not chunks:
        raise HTTPException(400, "Could not extract any text from the provided documents.")
    if len(chunks) > MAX_CHUNKS:
        chunks = chunks[:MAX_CHUNKS]
    embeddings = embed_texts(chunks)
    vectors = np.array(embeddings, dtype="float32")
    del embeddings
    gc.collect()
    index = faiss.IndexFlatL2(vectors.shape[1])
    index.add(vectors)
    del vectors
    gc.collect()
    # Store index, chunks AND document names
    SESSIONS[session_id] = (index, chunks, doc_names or [])
    # Initialise answer cache for this session
    ANSWER_CACHE[session_id] = {}
    return {"session_id": session_id, "chunks_created": len(chunks)}

# ── routes ────────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {
        "status": "ok",
        "service": "CybeSure SecureAnswer",
        "copyright": "© CybeSure Ltd. All rights reserved.",
        "data_residency": "UK",
        "gdpr_compliant": True,
        "data_retention": "none — in-memory processing only"
    }

@app.post("/upload/questionnaire")
async def upload_questionnaire(
    file: UploadFile = File(...),
    request: Request = None,
    db=Depends(get_db)
):
    data = await file.read()
    import base64
    original_b64 = base64.b64encode(data).decode()

    ext = file.filename.lower().rsplit(".", 1)[-1] if "." in file.filename else ""
    questions_text = extract_questions(file.filename, data)

    if not questions_text:
        raise HTTPException(
            400,
            f"No questions found in '{file.filename}'. "
            f"Check the file has a column headed 'Question'."
        )

    questions = [
        {"id": i, "text": q.strip(), "category": None}
        for i, q in enumerate(questions_text)
        if q.strip() and len(q.strip()) > 3
    ]

    # Extract full structure for Excel/XLS files (for matched export)
    structure = None
    if ext in ("xlsx", "xls"):
        try:
            structure = extract_questionnaire_structure(data, file.filename)
            # Store in global structure cache
            struct_key = hashlib.sha256(data[:512]).hexdigest()[:16]
            QUESTIONNAIRE_STRUCTURE[struct_key] = structure
        except Exception as e:
            print(f"Structure extraction error: {e}")
            struct_key = None
    else:
        struct_key = None

    return {
        "questions": questions,
        "total": len(questions),
        "filename": file.filename,
        "original_file": original_b64,
        "structure_key": struct_key,
        "sheet_count": len(structure["sheets"]) if structure else 1
    }

@app.post("/upload/documents")
async def upload_documents(
    files: List[UploadFile] = File(...),
    session_id: str = None,
    db=Depends(get_db)
):
    if not session_id:
        session_id = str(uuid.uuid4())
    raw, failed, doc_names = [], [], []
    for f in files:
        try:
            data = await f.read()
            parsed = parse_file(f.filename, data)
            if parsed:
                raw.extend(parsed)
                doc_names.append(f.filename)
            else:
                failed.append(f.filename)
            del data
            gc.collect()
        except Exception as e:
            print(f"Failed {f.filename}: {e}")
            failed.append(f.filename)
    if not raw:
        raise HTTPException(400, "Could not extract text from any documents.")
    result = build_index(raw, session_id, doc_names)
    result["files_processed"] = len(files) - len(failed)
    result["files_failed"] = failed
    result["document_names"] = doc_names
    return result

class UrlRequest(BaseModel):
    urls: List[str]
    session_id: Optional[str] = None

@app.post("/upload/documents-url")
async def upload_documents_url(req: UrlRequest):
    session_id = req.session_id or str(uuid.uuid4())
    raw, files_processed, files_failed, doc_names = [], 0, [], []
    for url in req.urls:
        try:
            file_pairs = fetch_all_from_url(url)
            if not file_pairs:
                files_failed.append(url)
                continue
            for fname, data in file_pairs:
                try:
                    parsed = parse_file(fname, data)
                    if parsed:
                        raw.extend(parsed)
                        files_processed += 1
                        doc_names.append(fname)
                    else:
                        files_failed.append(fname)
                except Exception as e:
                    print(f"Parse error {fname}: {e}")
                    files_failed.append(fname)
                del data
                gc.collect()
        except Exception as e:
            print(f"URL error {url}: {e}")
            files_failed.append(url)
    if not raw:
        raise HTTPException(
            400,
            "Could not fetch or read documents from the provided URLs. "
            "For Google Drive: share each file individually with 'Anyone with link can view' "
            "and paste the individual file share link."
        )
    result = build_index(raw, session_id, doc_names)
    result["files_processed"] = files_processed
    result["files_failed"] = files_failed
    result["documents_found"] = doc_names
    return result


# ── OneTrust / portal integration ────────────────────────────────────────────

# ── Portal Connector System ───────────────────────────────────────────────────
# Stores configured portal connections per session
PORTAL_CONNECTIONS: dict = {}  # connection_id -> connection config

SUPPORTED_PORTALS = {
    "onetrust": {
        "name": "OneTrust",
        "description": "OneTrust Privacy & Security Assessment Platform",
        "logo": "🛡️",
        "auth_type": "api_key",
        "base_url": "https://app.onetrust.com/api/assessment/v2",
        "docs": "https://developer.onetrust.com/onetrust/reference"
    },
    "vanta": {
        "name": "Vanta",
        "description": "Vanta Security & Compliance Platform",
        "logo": "✅",
        "auth_type": "api_key",
        "base_url": "https://api.vanta.com/v1",
        "docs": "https://developer.vanta.com/reference"
    },
    "whistic": {
        "name": "Whistic",
        "description": "Whistic Vendor Security Network",
        "logo": "🔍",
        "auth_type": "api_key",
        "base_url": "https://api.whistic.com/v1",
        "docs": "https://help.whistic.com/hc/en-us/articles/api"
    },
    "ecovadis": {
        "name": "EcoVadis",
        "description": "EcoVadis Sustainability Ratings",
        "logo": "🌿",
        "auth_type": "oauth2",
        "base_url": "https://api.ecovadis.com/v1",
        "docs": "https://developer.ecovadis.com"
    },
    "securityscorecard": {
        "name": "SecurityScorecard",
        "description": "SecurityScorecard Risk Assessment",
        "logo": "📊",
        "auth_type": "api_key",
        "base_url": "https://api.securityscorecard.io",
        "docs": "https://securityscorecard.com/docs/api"
    },
    "archer": {
        "name": "RSA Archer",
        "description": "RSA Archer GRC Platform",
        "logo": "🏹",
        "auth_type": "basic",
        "base_url": "",
        "docs": "https://community.rsa.com/docs"
    },
    "servicenow": {
        "name": "ServiceNow GRC",
        "description": "ServiceNow Governance Risk & Compliance",
        "logo": "⚙️",
        "auth_type": "basic",
        "base_url": "",
        "docs": "https://developer.servicenow.com"
    },
    "generic": {
        "name": "Custom / Generic Portal",
        "description": "Any portal with REST API or direct URL access",
        "logo": "🔗",
        "auth_type": "configurable",
        "base_url": "",
        "docs": ""
    }
}


class PortalConnectRequest(BaseModel):
    portal_type: str          # onetrust, vanta, whistic, ecovadis, generic etc
    portal_name: Optional[str] = ""      # custom name
    base_url: Optional[str] = ""         # custom base URL
    api_key: Optional[str] = ""          # API key / token
    client_id: Optional[str] = ""        # OAuth client ID
    client_secret: Optional[str] = ""    # OAuth client secret
    username: Optional[str] = ""         # Basic auth username
    password: Optional[str] = ""         # Basic auth password
    custom_headers: Optional[Dict] = {}  # Extra headers


class PortalRequest(BaseModel):
    portal_url: str
    session_id: str
    questions: List[dict] = []   # list of {id, text} dicts — avoids forward ref
    connection_id: Optional[str] = None


class PortalFetchRequest(BaseModel):
    connection_id: str
    session_id: str
    questionnaire_id: Optional[str] = ""   # specific questionnaire to fetch
    assessment_id: Optional[str] = ""


class PortalSubmitRequest(BaseModel):
    connection_id: str
    session_id: str
    results: List[dict]
    questionnaire_id: Optional[str] = ""
    assessment_id: Optional[str] = ""
    submit_mode: str = "draft"  # "draft" or "submit"


@app.get("/portal/connectors")
async def list_connectors():
    """List all supported portal connectors."""
    return {"connectors": SUPPORTED_PORTALS}


@app.post("/portal/connect")
async def portal_connect(req: PortalConnectRequest):
    """
    Configure and test a portal API connection.
    Returns a connection_id used for subsequent fetch/submit operations.
    """
    portal_info = SUPPORTED_PORTALS.get(req.portal_type, SUPPORTED_PORTALS["generic"])
    base_url = req.base_url or portal_info.get("base_url", "")

    # Build auth headers based on auth type
    headers = {"Content-Type": "application/json", "Accept": "application/json"}
    auth = None

    if req.api_key:
        # Try common API key header formats
        if req.portal_type == "onetrust":
            headers["Authorization"] = f"Bearer {req.api_key}"
        elif req.portal_type == "vanta":
            headers["Authorization"] = f"Bearer {req.api_key}"
        elif req.portal_type == "securityscorecard":
            headers["Token"] = req.api_key
        else:
            headers["Authorization"] = f"Bearer {req.api_key}"
            headers["X-API-Key"] = req.api_key

    if req.username and req.password:
        import base64
        creds = base64.b64encode(f"{req.username}:{req.password}".encode()).decode()
        headers["Authorization"] = f"Basic {creds}"

    if req.custom_headers:
        headers.update(req.custom_headers)

    # Test the connection
    connected = False
    test_message = "Connection not tested"
    test_endpoint = base_url

    if base_url:
        try:
            resp = requests.get(
                test_endpoint,
                headers=headers,
                timeout=10,
                allow_redirects=True
            )
            connected = resp.status_code in (200, 201, 204, 401, 403)
            if resp.status_code == 401:
                test_message = "Endpoint reachable but authentication failed — check API key"
            elif resp.status_code == 403:
                test_message = "Endpoint reachable but access denied — check permissions"
            elif resp.status_code == 200:
                test_message = "Connected successfully"
                connected = True
            else:
                test_message = f"HTTP {resp.status_code}"
        except Exception as e:
            test_message = f"Could not reach {base_url}: {str(e)[:100]}"
            connected = False

    # Store connection
    connection_id = hashlib.sha256(f"{req.portal_type}-{req.api_key or req.username}-{uuid.uuid4()}".encode()).hexdigest()[:20]
    PORTAL_CONNECTIONS[connection_id] = {
        "id": connection_id,
        "portal_type": req.portal_type,
        "portal_name": req.portal_name or portal_info["name"],
        "base_url": base_url,
        "headers": headers,
        "auth_type": portal_info.get("auth_type"),
        "connected": connected,
        "created_at": datetime.utcnow().isoformat()
    }

    return {
        "connection_id": connection_id,
        "portal": portal_info["name"],
        "connected": connected,
        "test_message": test_message,
        "base_url": base_url
    }


@app.post("/portal/fetch-questions")
async def portal_fetch_questions(req: PortalFetchRequest):
    """
    Fetch questionnaire questions from a connected portal.
    Handles OneTrust, Vanta, Whistic and generic REST APIs.
    """
    if req.connection_id not in PORTAL_CONNECTIONS:
        raise HTTPException(404, "Connection not found. Please connect to the portal first.")

    conn = PORTAL_CONNECTIONS[req.connection_id]
    base_url = conn["base_url"]
    headers = conn["headers"]
    portal_type = conn["portal_type"]

    questions = []
    raw_data = {}

    try:
        # Portal-specific question fetching
        if portal_type == "onetrust":
            # OneTrust Assessment API
            assessment_id = req.assessment_id or req.questionnaire_id
            if assessment_id:
                url = f"{base_url}/assessments/{assessment_id}/questions"
            else:
                url = f"{base_url}/assessments"
            resp = requests.get(url, headers=headers, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            raw_data = data
            # Extract questions from OneTrust format
            items = data.get("content", data.get("questions", data if isinstance(data, list) else []))
            for item in items:
                q_text = item.get("questionText") or item.get("name") or item.get("label") or ""
                q_id = item.get("questionId") or item.get("id") or ""
                if q_text and len(q_text) > 3:
                    questions.append({"id": q_id, "text": q_text, "category": item.get("category", "")})

        elif portal_type == "vanta":
            # Vanta API
            url = f"{base_url}/questionnaires/{req.questionnaire_id}/questions" if req.questionnaire_id else f"{base_url}/questionnaires"
            resp = requests.get(url, headers=headers, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            raw_data = data
            items = data.get("data", data.get("questions", []))
            for item in items:
                q_text = item.get("text") or item.get("question") or item.get("name") or ""
                if q_text and len(q_text) > 3:
                    questions.append({"id": item.get("id", ""), "text": q_text, "category": item.get("category", "")})

        elif portal_type == "whistic":
            url = f"{base_url}/questionnaires/{req.questionnaire_id}" if req.questionnaire_id else f"{base_url}/questionnaires"
            resp = requests.get(url, headers=headers, timeout=30)
            resp.raise_for_status()
            data = resp.json()
            raw_data = data
            items = data.get("questions", data.get("data", []))
            for item in items:
                q_text = item.get("text") or item.get("question") or ""
                if q_text and len(q_text) > 3:
                    questions.append({"id": item.get("id", ""), "text": q_text, "category": item.get("section", "")})

        else:
            # Generic — try common REST patterns
            url = base_url
            if req.questionnaire_id:
                url = f"{base_url}/{req.questionnaire_id}"
            resp = requests.get(url, headers=headers, timeout=30)
            resp.raise_for_status()

            # Try JSON
            try:
                data = resp.json()
                raw_data = data
                # Try multiple common structures
                items = (data.get("questions") or data.get("items") or
                        data.get("data") or data.get("content") or
                        (data if isinstance(data, list) else []))
                for i, item in enumerate(items):
                    if isinstance(item, str):
                        questions.append({"id": str(i), "text": item, "category": ""})
                    elif isinstance(item, dict):
                        q_text = (item.get("question") or item.get("text") or
                                item.get("name") or item.get("label") or
                                item.get("questionText") or "")
                        if q_text and len(q_text) > 3:
                            questions.append({"id": item.get("id", str(i)), "text": q_text, "category": item.get("category", item.get("section", ""))})
            except Exception:
                # HTML fallback
                soup = BeautifulSoup(resp.text, 'html.parser')
                for tag in soup.find_all(['p', 'li', 'td', 'label', 'h3', 'h4', 'th']):
                    text = tag.get_text(strip=True)
                    if text and len(text) > 10 and len(text) < 500:
                        questions.append({"id": str(len(questions)), "text": text, "category": ""})

    except requests.HTTPError as e:
        raise HTTPException(400, f"Portal API error: {e.response.status_code} — {e.response.text[:200]}")
    except Exception as e:
        raise HTTPException(400, f"Failed to fetch questions: {str(e)[:200]}")

    if not questions:
        return {
            "questions": [],
            "total": 0,
            "message": "No questions found. Check the questionnaire ID or try a different endpoint.",
            "raw_preview": str(raw_data)[:500]
        }

    return {
        "questions": questions,
        "total": len(questions),
        "portal": conn["portal_name"],
        "questionnaire_id": req.questionnaire_id
    }


@app.post("/portal/submit-answers")
async def portal_submit_answers(req: PortalSubmitRequest):
    """
    Submit AI-generated answers back to the portal.
    Supports draft save and final submission.
    """
    if req.connection_id not in PORTAL_CONNECTIONS:
        raise HTTPException(404, "Connection not found.")

    conn = PORTAL_CONNECTIONS[req.connection_id]
    base_url = conn["base_url"]
    headers = conn["headers"]
    portal_type = conn["portal_type"]

    submitted = 0
    failed = 0
    errors = []

    for result in req.results:
        q_id = result.get("portal_question_id") or result.get("question_id")
        answer = result.get("explanation") or result.get("answer", "")
        if not answer:
            continue

        try:
            if portal_type == "onetrust":
                url = f"{base_url}/assessments/{req.assessment_id}/questions/{q_id}/answer"
                payload = {"answer": answer, "status": "COMPLETE" if req.submit_mode == "submit" else "DRAFT"}
                resp = requests.put(url, json=payload, headers=headers, timeout=15)
                resp.raise_for_status()
                submitted += 1

            elif portal_type == "vanta":
                url = f"{base_url}/questionnaires/{req.questionnaire_id}/questions/{q_id}/response"
                payload = {"response": answer}
                resp = requests.post(url, json=payload, headers=headers, timeout=15)
                resp.raise_for_status()
                submitted += 1

            elif portal_type == "whistic":
                url = f"{base_url}/questionnaire-responses/{req.questionnaire_id}/questions/{q_id}"
                payload = {"answer": answer}
                resp = requests.patch(url, json=payload, headers=headers, timeout=15)
                resp.raise_for_status()
                submitted += 1

            else:
                # Generic PUT/POST — try common patterns
                url = f"{base_url}/{req.questionnaire_id}/answers/{q_id}" if req.questionnaire_id else f"{base_url}/answers/{q_id}"
                payload = {"answer": answer, "question_id": q_id}
                resp = requests.post(url, json=payload, headers=headers, timeout=15)
                if resp.status_code in (200, 201, 204):
                    submitted += 1
                else:
                    failed += 1
                    errors.append(f"Q{q_id}: HTTP {resp.status_code}")

        except Exception as e:
            failed += 1
            errors.append(f"Q{q_id}: {str(e)[:80]}")

    return {
        "submitted": submitted,
        "failed": failed,
        "total": len(req.results),
        "errors": errors[:10],
        "mode": req.submit_mode,
        "portal": conn["portal_name"]
    }


@app.get("/portal/connections")
async def list_connections():
    """List all active portal connections."""
    return {
        "connections": [
            {
                "id": c["id"],
                "portal_type": c["portal_type"],
                "portal_name": c["portal_name"],
                "connected": c["connected"],
                "created_at": c["created_at"]
            }
            for c in PORTAL_CONNECTIONS.values()
        ]
    }


@app.delete("/portal/connections/{connection_id}")
async def delete_connection(connection_id: str):
    """Remove a portal connection."""
    if connection_id in PORTAL_CONNECTIONS:
        del PORTAL_CONNECTIONS[connection_id]
    return {"message": "Connection removed"}

class Question(BaseModel):
    id: int
    text: str
    category: Optional[str] = None

class AnswerRequest(BaseModel):
    session_id: str
    questions: List[Question]
    is_demo: Optional[bool] = False
    org_id: Optional[int] = None
    industry_sector: Optional[str] = None

@app.post("/answer")
async def answer(
    req: AnswerRequest,
    db=Depends(get_db)
):
    if req.session_id not in SESSIONS:
        raise HTTPException(404, "Session not found. Upload documents first.")

    _, _, doc_names = SESSIONS[req.session_id]
    doc_names_str = ", ".join(doc_names) if doc_names else ""
    session_cache = ANSWER_CACHE.get(req.session_id, {})

    # Check subscription and get org tier
    is_demo = req.is_demo
    org = None
    if AUTH_ENABLED and db and req.org_id:
        try:
            org = db.query(Organisation).filter(Organisation.id == req.org_id).first()
            if org:
                is_demo = org.tier == "demo"
                is_test = org.tier == "test"
                if not is_test and org.remaining <= 0:
                    raise HTTPException(402, {
                        "error": "limit_reached",
                        "message": f"You have used all {org.total_limit} questionnaires in your subscription.",
                        "remaining": 0
                    })
        except HTTPException:
            raise
        except Exception as e:
            print(f"Org check error: {e}")

    # Demo mode: limit to DEMO_MAX_QUESTIONS total
    questions_to_process = req.questions
    demo_truncated = False
    if is_demo:
        max_q = DEMO_MAX_QUESTIONS
        if len(questions_to_process) > max_q:
            questions_to_process = questions_to_process[:max_q]
            demo_truncated = True

    results = []
    for q in questions_to_process:
        q_key = question_hash(q.text)

        # 1. Check session cache (fastest)
        if q.text in session_cache:
            r = session_cache[q.text].copy()
            r["question_id"] = q.id
            results.append(r)
            continue

        # 2. Check persistent question cache
        if q_key in QUESTION_CACHE:
            print(f"Cache hit for: {q.text[:60]}")
            r = QUESTION_CACHE[q_key].copy()
            r["question_id"] = q.id
            r["document_names"] = doc_names_str
            results.append(r)
            session_cache[q.text] = r
            continue

        # 3. Process with AI
        chunks = retrieve(req.session_id, q.text)
        r = ask_claude(q.text, chunks)

        result = {
            "question_id": q.id,
            "question": q.text,
            "category": q.category,
            "confidence": r.get("confidence", "Partial"),
            "confidence_pct": r.get("confidence_pct", 50),
            "explanation": str(r.get("explanation", "")),
            "improved_answer": str(r.get("improved_answer", "")),
            "improvement_notes": str(r.get("improvement_notes", "")),
            "sources": r.get("sources", []),
            "document_names": doc_names_str,
            "approval_status": "pending",
            "approved_by": None,
            "approval_token": str(uuid.uuid4())
        }

        # Store in both caches
        session_cache[q.text] = result
        QUESTION_CACHE[q_key] = result
        results.append(result)
        gc.collect()

    ANSWER_CACHE[req.session_id] = session_cache

    # Update usage counter and record analytics in DB
    remaining = None
    processing_secs = 0
    if AUTH_ENABLED and db and org:
        try:
            if org.tier not in ("test", "demo"):
                org.questionnaires_used += 1
            # Record run completion with timing
            run = db.query(QuestionnaireRun).filter(
                QuestionnaireRun.org_id == org.id,
                QuestionnaireRun.status == "processing"
            ).order_by(QuestionnaireRun.created_at.desc()).first()
            if run:
                now = datetime.utcnow()
                processing_secs = (now - run.created_at).total_seconds()
                run.status = "complete"
                run.completed_at = now
                run.question_count = len(results)
                run.processing_secs = processing_secs
                if hasattr(req, 'industry_sector') and req.industry_sector:
                    run.industry_sector = req.industry_sector
            db.commit()
            remaining = org.remaining
        except Exception as e:
            print(f"Usage/analytics update error: {e}")

    return {
        "results": results,
        "total": len(results),
        "demo_truncated": demo_truncated,
        "demo_max": DEMO_MAX_QUESTIONS if is_demo else None,
        "remaining": remaining,
        "processing_secs": processing_secs
    }


# ── Approval workflow ─────────────────────────────────────────────────────────

# Approver/Reviewer directory: org_id -> list of contacts
APPROVER_DIRECTORY: dict = {}

class ApproverContact(BaseModel):
    name: str
    email: str
    role: str = "approver"  # "approver" or "reviewer"
    title: Optional[str] = ""  # e.g. "CISO", "Head of IT", "DPO"

class ApproverContactResponse(BaseModel):
    id: str
    name: str
    email: str
    role: str
    title: Optional[str] = ""

@app.get("/admin/approvers")
async def list_approvers(db=Depends(get_db)):
    """List all configured approvers and reviewers for this organisation."""
    # Get org from token if available
    org_id = "default"
    try:
        from fastapi import Request
    except Exception:
        pass
    contacts = APPROVER_DIRECTORY.get(org_id, [])
    # Also include registered users with approver/contributor roles
    if AUTH_ENABLED and db:
        try:
            users = db.query(User).filter(
                User.role.in_(["approver", "contributor", "admin"]),
                User.is_active == True
            ).all()
            existing_emails = {c["email"] for c in contacts}
            for u in users:
                if u.email not in existing_emails:
                    contacts.append({
                        "id": f"user-{u.id}",
                        "name": u.full_name,
                        "email": u.email,
                        "role": u.role,
                        "title": u.role.title()
                    })
        except Exception as e:
            print(f"User lookup error: {e}")
    return {"contacts": contacts}

@app.post("/admin/approvers")
async def add_approver(contact: ApproverContact, db=Depends(get_db)):
    """Add an approver or reviewer to the directory."""
    org_id = "default"
    if org_id not in APPROVER_DIRECTORY:
        APPROVER_DIRECTORY[org_id] = []
    # Check for duplicate
    for c in APPROVER_DIRECTORY[org_id]:
        if c["email"].lower() == contact.email.lower():
            raise HTTPException(400, f"{contact.email} is already in the directory")
    contact_id = hashlib.sha256(contact.email.encode()).hexdigest()[:8]
    entry = {
        "id": contact_id,
        "name": contact.name,
        "email": contact.email,
        "role": contact.role,
        "title": contact.title or ""
    }
    APPROVER_DIRECTORY[org_id].append(entry)
    return {"message": f"{contact.name} added as {contact.role}", "contact": entry}

@app.delete("/admin/approvers/{contact_id}")
async def remove_approver(contact_id: str):
    """Remove an approver or reviewer from the directory."""
    org_id = "default"
    contacts = APPROVER_DIRECTORY.get(org_id, [])
    APPROVER_DIRECTORY[org_id] = [c for c in contacts if c["id"] != contact_id]
    return {"message": "Contact removed"}

def send_email(to: str, subject: str, html_body: str) -> bool:
    """Send email via SMTP. Returns True if sent successfully."""
    if not SMTP_USER or not SMTP_PASS:
        print("Email not configured — SMTP_USER/SMTP_PASS env vars missing")
        return False
    try:
        msg = MIMEMultipart("alternative")
        msg["From"] = SMTP_USER
        msg["To"] = to
        msg["Subject"] = subject
        msg.attach(MIMEText(html_body, "html"))
        with smtplib.SMTP(SMTP_HOST, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASS)
            server.sendmail(SMTP_USER, to, msg.as_string())
        return True
    except Exception as e:
        print(f"Email error: {e}")
        return False


class ApprovalRequest(BaseModel):
    session_id: str
    question_id: int
    question: str
    answer: str
    improved_answer: str
    approver_email: str
    approver_name: str
    approver_type: str = "approver"  # "approver" or "reviewer"
    reviewer_email: Optional[str] = None   # optional second reviewer
    reviewer_name: Optional[str] = None
    requester_name: Optional[str] = "CybeSure SecureAnswer"


class ApprovalDecision(BaseModel):
    token: str
    decision: str  # "approved" or "rejected"
    comments: Optional[str] = ""


@app.post("/approval/request")
async def request_approval(req: ApprovalRequest, background_tasks: BackgroundTasks):
    """Send approval/review request emails."""
    token = hashlib.sha256(f"{req.session_id}-{req.question_id}-{uuid.uuid4()}".encode()).hexdigest()[:32]

    is_review = req.approver_type == "reviewer"
    action_label = "Review" if is_review else "Approval"

    APPROVALS[token] = {
        "token": token,
        "session_id": req.session_id,
        "question_id": req.question_id,
        "question": req.question,
        "answer": req.answer,
        "improved_answer": req.improved_answer,
        "approver_email": req.approver_email,
        "approver_name": req.approver_name,
        "approver_type": req.approver_type,
        "reviewer_email": req.reviewer_email,
        "reviewer_name": req.reviewer_name,
        "requester_name": req.requester_name,
        "status": "pending",
        "review_status": "pending",
        "comments": "",
        "requested_at": datetime.utcnow().isoformat(),
        "decided_at": None
    }

    approve_url = f"{APP_URL}/approval/decide?token={token}&decision=approved"
    reject_url  = f"{APP_URL}/approval/decide?token={token}&decision=rejected"
    review_url  = f"{APP_URL}/approval/review/{token}"

    def build_email(recipient_name: str, recipient_type: str) -> str:
        action_buttons = "" if recipient_type == "reviewer" else f"""
        <div style="margin-bottom:24px">
          <a href="{approve_url}" style="display:inline-block;background:#00c8e0;color:#0f2035;padding:12px 28px;text-decoration:none;font-weight:700;font-size:14px;border-radius:3px;margin-right:12px">✓ Approve Answer</a>
          <a href="{reject_url}" style="display:inline-block;background:#fff;color:#e84855;padding:12px 28px;text-decoration:none;font-weight:700;font-size:14px;border-radius:3px;border:2px solid #e84855">✗ Reject / Request Changes</a>
        </div>"""
        review_only = f"""
        <div style="margin-bottom:24px">
          <a href="{review_url}" style="display:inline-block;background:#00c8e0;color:#0f2035;padding:12px 28px;text-decoration:none;font-weight:700;font-size:14px;border-radius:3px">Add Review Comments</a>
        </div>""" if recipient_type == "reviewer" else ""

        role_label = "review" if recipient_type == "reviewer" else "approve or reject"
        return f"""
        <div style="font-family:Arial,sans-serif;max-width:680px;margin:0 auto">
          <div style="background:#0f2035;padding:24px 32px;border-bottom:3px solid #00c8e0">
            <h1 style="color:#fff;font-size:20px;margin:0">Cybe<span style="color:#00c8e0">Sure</span> SecureAnswer</h1>
            <p style="color:#7a9cbf;font-size:13px;margin:4px 0 0">AI Security Questionnaire System</p>
          </div>
          <div style="padding:32px;background:#f9fafb;border:1px solid #e2e8f0">
            <h2 style="color:#0f2035;font-size:18px;margin:0 0 8px">Answer {action_label} Request</h2>
            <p style="color:#4a6a8a;font-size:14px;margin:0 0 24px">
              Hi <strong>{recipient_name}</strong>, <strong>{req.requester_name}</strong> is requesting you to <strong>{role_label}</strong> the following questionnaire answer.
            </p>
            <div style="background:#fff;border:1px solid #e2e8f0;border-radius:4px;padding:20px;margin-bottom:20px">
              <p style="font-size:12px;color:#4a6a8a;font-weight:700;text-transform:uppercase;letter-spacing:1px;margin:0 0 8px">Question</p>
              <p style="color:#0f2035;font-size:15px;font-weight:600;margin:0 0 20px">{req.question}</p>
              <p style="font-size:12px;color:#4a6a8a;font-weight:700;text-transform:uppercase;letter-spacing:1px;margin:0 0 8px">Proposed Answer</p>
              <p style="color:#334155;font-size:14px;line-height:1.7;margin:0 0 20px;padding:12px;background:#f8fafc;border-left:3px solid #00c8e0">{req.answer}</p>
              {"<p style='font-size:12px;color:#4a6a8a;font-weight:700;text-transform:uppercase;margin:0 0 8px'>Enhanced Answer (95%+)</p><p style='color:#334155;font-size:14px;line-height:1.7;padding:12px;background:#f0fdf4;border-left:3px solid #00d4a0;margin:0 0 20px'>" + req.improved_answer + "</p>" if req.improved_answer else ""}
            </div>
            {action_buttons}{review_only}
            <p style="color:#7a9cbf;font-size:12px">Or <a href="{review_url}" style="color:#00c8e0">view full details and add comments</a> before deciding.</p>
          </div>
          <div style="padding:16px 32px;background:#0f2035;text-align:center">
            <p style="color:#4a6a8a;font-size:11px;margin:0">CybeSure SecureAnswer — AI Compliance Engine | © CybeSure Ltd</p>
          </div>
        </div>"""

    # Send to primary approver/reviewer
    background_tasks.add_task(
        send_email,
        req.approver_email,
        f"Answer {action_label} Required: {req.question[:60]}...",
        build_email(req.approver_name, req.approver_type)
    )

    # Send to secondary reviewer if provided
    if req.reviewer_email and req.reviewer_name:
        background_tasks.add_task(
            send_email,
            req.reviewer_email,
            f"Answer Review Required: {req.question[:60]}...",
            build_email(req.reviewer_name, "reviewer")
        )

    recipients = [req.approver_name]
    if req.reviewer_name:
        recipients.append(req.reviewer_name)

    return {
        "token": token,
        "status": "pending",
        "message": f"{action_label} request sent to {', '.join(recipients)}",
        "review_url": review_url
    }


@app.get("/approval/review/{token}", response_class=HTMLResponse)
async def review_approval(token: str):
    """Show approval review page for approver."""
    if token not in APPROVALS:
        return HTMLResponse("<h2>Approval request not found or expired.</h2>", status_code=404)
    
    rec = APPROVALS[token]
    approve_url = f"{APP_URL}/approval/decide?token={token}&decision=approved"
    reject_url  = f"{APP_URL}/approval/decide?token={token}&decision=rejected"
    
    status_colour = {"approved": "#00d4a0", "rejected": "#e84855", "pending": "#f5a623"}.get(rec["status"], "#f5a623")
    
    return HTMLResponse(f"""
    <!DOCTYPE html>
    <html>
    <head>
      <title>CybeSure — Answer Approval</title>
      <link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@600;700&family=Open+Sans:wght@400;500&display=swap" rel="stylesheet">
      <style>
        body {{font-family:'Open Sans',sans-serif;background:#0b1829;color:#dce8f5;margin:0;padding:0}}
        .header {{background:#0f2035;border-bottom:3px solid #00c8e0;padding:20px 40px;display:flex;align-items:center;gap:12px}}
        .logo {{font-family:'Montserrat',sans-serif;font-size:22px;font-weight:700;color:#fff}}
        .logo em {{color:#00c8e0;font-style:normal}}
        .container {{max-width:800px;margin:40px auto;padding:0 24px}}
        .card {{background:#0f2035;border:1px solid #1e3a5a;padding:28px;margin-bottom:20px;border-radius:4px}}
        .label {{font-size:11px;font-weight:700;color:#7a9cbf;letter-spacing:1px;text-transform:uppercase;margin-bottom:8px}}
        .question {{font-size:18px;font-weight:600;color:#fff;line-height:1.5;margin-bottom:24px}}
        .answer-box {{background:#152a42;border-left:3px solid #00c8e0;padding:16px;font-size:14px;line-height:1.8;color:#dce8f5;border-radius:2px;margin-bottom:16px}}
        .improved-box {{background:#0d2a1f;border-left:3px solid #00d4a0;padding:16px;font-size:14px;line-height:1.8;color:#dce8f5;border-radius:2px}}
        .status-badge {{display:inline-block;padding:4px 14px;border-radius:2px;font-size:12px;font-weight:700;background:{status_colour}22;color:{status_colour};border:1px solid {status_colour}55;margin-bottom:20px}}
        .btn-approve {{display:inline-block;background:#00c8e0;color:#0f2035;padding:12px 32px;text-decoration:none;font-weight:700;font-size:14px;border-radius:3px;margin-right:12px;font-family:'Montserrat',sans-serif}}
        .btn-reject {{display:inline-block;background:transparent;color:#e84855;padding:12px 32px;text-decoration:none;font-weight:700;font-size:14px;border-radius:3px;border:2px solid #e84855;font-family:'Montserrat',sans-serif}}
        .comments {{width:100%;background:#152a42;border:1px solid #1e3a5a;color:#dce8f5;padding:12px;font-size:14px;border-radius:3px;margin-bottom:16px;min-height:80px;font-family:'Open Sans',sans-serif}}
        .decided {{background:#1a2f4a;border:1px solid #2d5a80;padding:20px;border-radius:4px;text-align:center}}
      </style>
    </head>
    <body>
    <div class="header">
      <div class="logo">Cybe<em>Sure</em></div>
      <span style="color:#7a9cbf;font-size:14px;margin-left:8px">Answer Approval Review</span>
    </div>
    <div class="container">
      <div class="card">
        <div class="status-badge">{rec['status'].upper()}</div>
        <div class="label">Question</div>
        <div class="question">{rec['question']}</div>
        <div class="label">Proposed Answer</div>
        <div class="answer-box">{rec['answer']}</div>
        {"<div class='label'>Enhanced Answer (95%+ Score)</div><div class='improved-box'>" + rec['improved_answer'] + "</div>" if rec.get('improved_answer') else ""}
      </div>
      {"<div class='decided'><p style='font-size:18px;font-weight:600;color:" + status_colour + "'>" + rec['status'].upper() + "</p><p style='color:#7a9cbf'>" + (rec.get('comments') or 'No comments') + "</p></div>" if rec['status'] != 'pending' else f"""
      <div class="card">
        <div class="label">Your Decision</div>
        <textarea class="comments" id="comments" placeholder="Add any comments or requested changes (optional)..."></textarea>
        <a href="{approve_url}" class="btn-approve" onclick="addComments(this,'approved')">✓ Approve</a>
        <a href="{reject_url}" class="btn-reject" onclick="addComments(this,'rejected')">✗ Reject / Request Changes</a>
      </div>
      <script>
        function addComments(el, decision) {{
          const c = document.getElementById('comments').value;
          el.href = `{APP_URL}/approval/decide?token={token}&decision=${{decision}}&comments=${{encodeURIComponent(c)}}`;
        }}
      </script>
      """}
    </div>
    </body>
    </html>
    """)


@app.get("/approval/decide")
async def decide_approval(token: str, decision: str, comments: str = "", background_tasks: BackgroundTasks = None):
    """Process approval decision and notify requester."""
    if token not in APPROVALS:
        return HTMLResponse("<h2>Approval request not found or expired.</h2>", status_code=404)
    
    if decision not in ("approved", "rejected"):
        raise HTTPException(400, "Decision must be 'approved' or 'rejected'")
    
    rec = APPROVALS[token]
    rec["status"] = decision
    rec["comments"] = comments
    rec["decided_at"] = datetime.utcnow().isoformat()

    status_colour = "#00d4a0" if decision == "approved" else "#e84855"
    status_word = "Approved" if decision == "approved" else "Rejected"

    # Send notification back to requester
    notify_html = f"""
    <div style="font-family:Arial,sans-serif;max-width:680px;margin:0 auto">
      <div style="background:#0f2035;padding:24px 32px;border-bottom:3px solid #00c8e0">
        <h1 style="color:#fff;font-size:20px;margin:0">Cybe<span style="color:#00c8e0">Sure</span> SecureAnswer</h1>
      </div>
      <div style="padding:32px;background:#f9fafb;border:1px solid #e2e8f0">
        <h2 style="color:{status_colour};font-size:20px;margin:0 0 8px">Answer {status_word}</h2>
        <p style="color:#4a6a8a;font-size:14px;margin:0 0 20px">
          <strong>{rec['approver_name']}</strong> has <strong>{decision}</strong> the answer to:
        </p>
        <p style="color:#0f2035;font-weight:600;font-size:15px;margin:0 0 16px">{rec['question']}</p>
        {f"<div style='background:#fff3cd;border:1px solid #ffc107;padding:12px 16px;border-radius:3px;margin-bottom:16px'><p style='font-weight:700;color:#856404;margin:0 0 4px'>Comments from approver:</p><p style='color:#856404;margin:0'>{comments}</p></div>" if comments else ""}
        <p style="color:#4a6a8a;font-size:13px">Log in to CybeSure SecureAnswer to view and update this answer.</p>
      </div>
    </div>
    """

    if background_tasks:
        background_tasks.add_task(
            send_email,
            SMTP_USER,  # notify the system/requester
            f"Answer {status_word}: {rec['question'][:60]}...",
            notify_html
        )

    return HTMLResponse(f"""
    <!DOCTYPE html>
    <html>
    <head>
      <title>CybeSure — Decision Recorded</title>
      <link href="https://fonts.googleapis.com/css2?family=Montserrat:wght@600;700&family=Open+Sans:wght@400;500&display=swap" rel="stylesheet">
      <style>
        body {{font-family:'Open Sans',sans-serif;background:#0b1829;color:#dce8f5;margin:0;display:flex;align-items:center;justify-content:center;min-height:100vh}}
        .box {{background:#0f2035;border:1px solid #1e3a5a;padding:48px;border-radius:4px;text-align:center;max-width:480px}}
        .logo {{font-family:'Montserrat',sans-serif;font-size:22px;font-weight:700;color:#fff;margin-bottom:24px}}
        .logo em {{color:#00c8e0;font-style:normal}}
        .status {{font-size:48px;margin-bottom:16px}}
        h2 {{color:{status_colour};font-family:'Montserrat',sans-serif;margin:0 0 12px}}
        p {{color:#7a9cbf;font-size:14px;line-height:1.6}}
      </style>
    </head>
    <body>
      <div class="box">
        <div class="logo">Cybe<em>Sure</em></div>
        <div class="status">{"✅" if decision == "approved" else "❌"}</div>
        <h2>{status_word}</h2>
        <p>Your decision has been recorded{f' with comments: <em>{comments}</em>' if comments else ''}.</p>
        <p style="margin-top:16px;color:#4a6a8a;font-size:12px">You can close this window.</p>
      </div>
    </body>
    </html>
    """)


@app.get("/approval/status/{token}")
async def approval_status(token: str):
    """Check the status of an approval request."""
    if token not in APPROVALS:
        raise HTTPException(404, "Approval token not found")
    rec = APPROVALS[token]
    return {
        "token": token,
        "question": rec["question"],
        "status": rec["status"],
        "approver": rec["approver_name"],
        "comments": rec.get("comments", ""),
        "requested_at": rec["requested_at"],
        "decided_at": rec.get("decided_at")
    }

class ResultItem(BaseModel):
    question_id: int
    question: str
    category: Optional[str] = None
    confidence: str
    confidence_pct: Optional[int] = 50
    explanation: str = ""
    improved_answer: Optional[str] = ""
    improvement_notes: Optional[str] = ""
    sources: List[str] = []
    document_names: Optional[str] = ""
    approval_status: Optional[str] = "pending"
    approved_by: Optional[str] = None
    approval_token: Optional[str] = None

    def dict(self, **kwargs):
        d = super().dict(**kwargs)
        # Ensure these are always strings
        for field in ["explanation","improved_answer","improvement_notes","document_names"]:
            if isinstance(d.get(field), list):
                d[field] = "\n".join(str(x) for x in d[field])
            elif d.get(field) is None:
                d[field] = ""
        return d

class ExportRequest(BaseModel):
    results: List[ResultItem]
    format: str
    original_file: Optional[str] = None
    structure_key: Optional[str] = None  # key to look up full questionnaire structure

@app.post("/export")
async def export(req: ExportRequest):
    import base64
    from openpyxl import load_workbook
    from openpyxl.styles import Alignment, PatternFill, Font, Border, Side
    from openpyxl.utils import get_column_letter

    fmt = req.format.lower()
    results = req.results

    # ── Build approval lookup: question text → approved_by name ──────────────
    approval_lookup: dict = {}
    for token, rec in APPROVALS.items():
        if rec.get("status") == "approved":
            approval_lookup[rec["question"].strip()] = rec.get("approver_name", "")

    # ── Helper: get approval name for a result ────────────────────────────────
    def get_approver(r) -> str:
        name = approval_lookup.get(r.question.strip(), "")
        if not name and r.approved_by:
            name = r.approved_by
        return name

    # ── JSON ─────────────────────────────────────────────────────────────────
    if fmt == "json":
        content = json.dumps([r.dict() for r in results], indent=2)
        return StreamingResponse(io.BytesIO(content.encode()),
            media_type="application/json",
            headers={"Content-Disposition": "attachment; filename=secureanswer_results.json"})

    # ── EXCEL ─────────────────────────────────────────────────────────────────
    if fmt == "excel":
        result_map = {r.question.strip(): r for r in results}

        # ── FORMAT 1: Matched original file — preserve ALL columns, ALL sheets ──
        structure = QUESTIONNAIRE_STRUCTURE.get(req.structure_key) if req.structure_key else None

        if structure and structure.get("sheets"):
            try:
                from openpyxl import Workbook as NewWB
                from openpyxl.styles import Alignment, PatternFill, Font
                from openpyxl.utils import get_column_letter

                wb_out = NewWB()
                wb_out.remove(wb_out.active)  # remove default sheet

                for sheet_info in structure["sheets"]:
                    ws = wb_out.create_sheet(title=sheet_info["name"])
                    cols = sheet_info["columns"]
                    q_col = sheet_info["question_col"]
                    a_col = sheet_info["answer_col"]

                    # Determine answer column index
                    if a_col and a_col in cols:
                        a_col_idx = cols.index(a_col)
                    else:
                        # Add Answer column after last existing column
                        a_col = "Answer"
                        cols = cols + ["Answer"]
                        a_col_idx = len(cols) - 1

                    # Add new output columns at end
                    conf_col_name   = "Confidence"
                    pct_col_name    = "Confidence %"
                    src_col_name    = "Policy Sources"
                    appr_col_name   = "Approved By"
                    out_cols = cols + [conf_col_name, pct_col_name, src_col_name, appr_col_name]

                    # Write header row
                    hdr_fill = PatternFill("solid", fgColor="0F2035")
                    hdr_font = Font(bold=True, color="FFFFFF", name="Calibri")
                    for ci, col_name in enumerate(out_cols, 1):
                        cell = ws.cell(row=1, column=ci, value=col_name)
                        cell.fill = hdr_fill
                        cell.font = hdr_font
                        cell.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

                    # Write data rows — preserve all original values, only fill Answer and new cols
                    for ri, row_info in enumerate(sheet_info["rows"], 2):
                        row_data = row_info["row_data"]
                        question = row_info["question"]
                        result = result_map.get(question)

                        for ci, col_name in enumerate(out_cols, 1):
                            cell = ws.cell(row=ri, column=ci)
                            cell.alignment = Alignment(wrap_text=True, vertical='top')

                            if col_name in (conf_col_name, pct_col_name, src_col_name, appr_col_name):
                                # New output columns
                                if result:
                                    if col_name == conf_col_name:
                                        cell.value = result.confidence
                                        if result.confidence == "Yes":
                                            cell.fill = PatternFill("solid", fgColor="C6EFCE")
                                            cell.font = Font(color="276221", bold=True)
                                        elif result.confidence == "No":
                                            cell.fill = PatternFill("solid", fgColor="FFC7CE")
                                            cell.font = Font(color="9C0006", bold=True)
                                        else:
                                            cell.fill = PatternFill("solid", fgColor="FFEB9C")
                                            cell.font = Font(color="7D6608", bold=True)
                                    elif col_name == pct_col_name:
                                        cell.value = f"{result.confidence_pct}%"
                                    elif col_name == src_col_name:
                                        src = result.document_names or ""
                                        src = src.replace(", ", "\n").replace(" | ", "\n").strip()
                                        cell.value = src
                                    elif col_name == appr_col_name:
                                        cell.value = get_approver(result)
                            elif col_name == a_col:
                                # Answer column — fill with AI answer
                                cell.value = result.explanation if result else row_data.get(col_name, "")
                                cell.alignment = Alignment(wrap_text=True, vertical='top')
                            else:
                                # All other original columns — copy exactly as-is
                                orig_val = row_data.get(col_name, "")
                                cell.value = orig_val if orig_val and orig_val.lower() not in ('nan','none') else ""

                    # Set column widths
                    for ci, col_name in enumerate(out_cols, 1):
                        if col_name == a_col or col_name == "Answer":
                            ws.column_dimensions[get_column_letter(ci)].width = 80
                        elif col_name == src_col_name:
                            ws.column_dimensions[get_column_letter(ci)].width = 35
                        elif col_name in (conf_col_name, pct_col_name):
                            ws.column_dimensions[get_column_letter(ci)].width = 14
                        elif col_name == appr_col_name:
                            ws.column_dimensions[get_column_letter(ci)].width = 22
                        else:
                            ws.column_dimensions[get_column_letter(ci)].width = 18

                    ws.freeze_panes = "A2"

                buf1 = io.BytesIO()
                wb_out.save(buf1)
                buf1.seek(0)
                return StreamingResponse(buf1,
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": "attachment; filename=SecureAnswer_Completed.xlsx"})

            except Exception as e:
                print(f"Structure export error: {e}")
                import traceback; traceback.print_exc()

        # ── FORMAT 1b: Fallback — use original file bytes if no structure ─────
        elif req.original_file:
            try:
                import base64
                from openpyxl import load_workbook
                from openpyxl.styles import Alignment, PatternFill, Font
                from openpyxl.utils import get_column_letter

                original_bytes = base64.b64decode(req.original_file)
                wb = load_workbook(io.BytesIO(original_bytes))
                ws = wb.active

                headers = {}
                max_existing_col = ws.max_column
                for col in range(1, max_existing_col + 1):
                    cell = ws.cell(row=1, column=col)
                    if cell.value:
                        headers[str(cell.value).strip().lower()] = col

                q_col = next((headers[k] for k in ('question','questions','requirement','control') if k in headers), 2)
                ans_col = next((headers[k] for k in ('answer','answers','response') if k in headers), None)
                if ans_col is None:
                    ans_col = max_existing_col + 1
                    ws.cell(row=1, column=ans_col, value="Answer").font = Font(bold=True)

                next_col = max(max_existing_col, ans_col) + 1
                conf_col, pct_col, src_col, appr_col = next_col, next_col+1, next_col+2, next_col+3
                bold = Font(bold=True)
                for col, name in [(conf_col,"Confidence"),(pct_col,"Confidence %"),(src_col,"Policy Sources"),(appr_col,"Approved By")]:
                    ws.cell(row=1, column=col, value=name).font = bold

                for row in range(2, ws.max_row + 1):
                    q_text = str(ws.cell(row=row, column=q_col).value or "").strip()
                    r = result_map.get(q_text)
                    if not r: continue
                    ac = ws.cell(row=row, column=ans_col)
                    ac.value = r.explanation
                    ac.alignment = Alignment(wrap_text=True, vertical='top')
                    cc = ws.cell(row=row, column=conf_col)
                    cc.value = r.confidence
                    if r.confidence == "Yes": cc.fill = PatternFill("solid", fgColor="C6EFCE"); cc.font = Font(color="276221", bold=True)
                    elif r.confidence == "No": cc.fill = PatternFill("solid", fgColor="FFC7CE"); cc.font = Font(color="9C0006", bold=True)
                    else: cc.fill = PatternFill("solid", fgColor="FFEB9C"); cc.font = Font(color="7D6608", bold=True)
                    ws.cell(row=row, column=pct_col).value = f"{r.confidence_pct}%"
                    src = (r.document_names or "").replace(", ","\n").replace(" | ","\n").strip()
                    sc = ws.cell(row=row, column=src_col)
                    sc.value = src
                    sc.alignment = Alignment(wrap_text=True, vertical='top')
                    ws.cell(row=row, column=appr_col).value = get_approver(r)

                for col, w in [(ans_col,80),(conf_col,14),(pct_col,14),(src_col,35),(appr_col,22)]:
                    ws.column_dimensions[get_column_letter(col)].width = w

                buf1 = io.BytesIO()
                wb.save(buf1)
                buf1.seek(0)
                return StreamingResponse(buf1,
                    media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    headers={"Content-Disposition": "attachment; filename=SecureAnswer_Completed.xlsx"})
            except Exception as e:
                print(f"Fallback export error: {e}")
                ws = wb.active

                # Find all header columns
                headers = {}
                max_existing_col = ws.max_column
                for col in range(1, max_existing_col + 1):
                    cell = ws.cell(row=1, column=col)
                    if cell.value:
                        headers[str(cell.value).strip().lower()] = col

                # Find question column
                q_col = None
                for key in ('question', 'questions', 'requirement', 'control'):
                    if key in headers:
                        q_col = headers[key]
                        break
                if q_col is None:
                    q_col = 2  # default column B

                # Find or create Answer column
                ans_col = None
                for key in ('answer', 'answers', 'response'):
                    if key in headers:
                        ans_col = headers[key]
                        break
                if ans_col is None:
                    ans_col = max_existing_col + 1
                    ws.cell(row=1, column=ans_col, value="Answer").font = Font(bold=True)

                # Add new columns AFTER all existing columns (don't touch existing ones)
                next_col = max(max_existing_col, ans_col) + 1

                conf_col = next_col;     next_col += 1
                pct_col  = next_col;     next_col += 1
                src_col  = next_col;     next_col += 1
                appr_col = next_col

                # Write new headers
                bold = Font(bold=True)
                ws.cell(row=1, column=conf_col, value="Confidence").font = bold
                ws.cell(row=1, column=pct_col,  value="Confidence %").font = bold
                ws.cell(row=1, column=src_col,  value="Policy Sources").font = bold
                ws.cell(row=1, column=appr_col, value="Approved By").font = bold

                # Fill answers row by row — ONLY touch question and new columns
                for row in range(2, ws.max_row + 1):
                    q_cell = ws.cell(row=row, column=q_col)
                    if not q_cell.value:
                        continue
                    q_text = str(q_cell.value).strip()
                    r = result_map.get(q_text)
                    if not r:
                        continue

                    # Answer column
                    ac = ws.cell(row=row, column=ans_col)
                    ac.value = r.explanation
                    ac.alignment = Alignment(wrap_text=True, vertical='top')

                    # Confidence
                    cc = ws.cell(row=row, column=conf_col)
                    cc.value = r.confidence
                    if r.confidence == "Yes":
                        cc.fill = PatternFill("solid", fgColor="C6EFCE")
                        cc.font = Font(color="276221", bold=True)
                    elif r.confidence == "No":
                        cc.fill = PatternFill("solid", fgColor="FFC7CE")
                        cc.font = Font(color="9C0006", bold=True)
                    else:
                        cc.fill = PatternFill("solid", fgColor="FFEB9C")
                        cc.font = Font(color="7D6608", bold=True)

                    # Confidence %
                    ws.cell(row=row, column=pct_col).value = f"{r.confidence_pct}%"

        # ── FORMAT 2: Clean summary report (always available as fallback) ──────
        from openpyxl import Workbook as NewWB
        from openpyxl.styles import PatternFill, Font, Alignment, Border, Side

        wb2 = NewWB()
        ws2 = wb2.active
        ws2.title = "SecureAnswer Results"

        # Header row
        headers2 = ["Ref", "Question", "Answer", "Confidence", "Confidence %",
                    "Policy Sources", "Approved By"]
        hdr_fill = PatternFill("solid", fgColor="0F2035")
        hdr_font = Font(bold=True, color="FFFFFF", name="Calibri")
        for ci, h in enumerate(headers2, 1):
            cell = ws2.cell(row=1, column=ci, value=h)
            cell.fill = hdr_fill
            cell.font = hdr_font
            cell.alignment = Alignment(horizontal='center', vertical='center')

        # Data rows
        for r in results:
            row_data = [
                r.question_id + 1,
                r.question,
                r.explanation,
                r.confidence,
                f"{r.confidence_pct}%",
                r.document_names.replace(", ", "\n").replace(" | ", "\n").strip() if r.document_names else "\n".join(str(s) for s in r.sources),
                get_approver(r)
            ]
            ri = r.question_id + 2
            for ci, val in enumerate(row_data, 1):
                cell = ws2.cell(row=ri, column=ci, value=val)
                cell.alignment = Alignment(wrap_text=True, vertical='top')

            # Colour confidence cell
            cc = ws2.cell(row=ri, column=4)
            if r.confidence == "Yes":
                cc.fill = PatternFill("solid", fgColor="C6EFCE")
                cc.font = Font(color="276221", bold=True)
            elif r.confidence == "No":
                cc.fill = PatternFill("solid", fgColor="FFC7CE")
                cc.font = Font(color="9C0006", bold=True)
            else:
                cc.fill = PatternFill("solid", fgColor="FFEB9C")
                cc.font = Font(color="7D6608", bold=True)

        # Column widths
        widths = [6, 40, 80, 14, 14, 40, 22]
        for ci, w in enumerate(widths, 1):
            ws2.column_dimensions[get_column_letter(ci)].width = w

        ws2.freeze_panes = "A2"
        ws2.row_dimensions[1].height = 20

        buf2 = io.BytesIO()
        wb2.save(buf2)
        buf2.seek(0)
        return StreamingResponse(buf2,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": "attachment; filename=SecureAnswer_Report.xlsx"})

    if fmt == "word":
        doc = Document()
        doc.add_heading("CybeSure SecureAnswer — Compliance Results", 0)
        p = doc.add_paragraph()
        p.add_run(f"© {datetime.utcnow().year} CybeSure Ltd. All rights reserved. SecureAnswer™ is a trademark of CybeSure Ltd.").italic = True
        p.runs[0].font.size = 91440
        p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
        doc.add_paragraph("")
        for r in results:
            doc.add_heading(f"Q{r.question_id+1}: {r.question[:120]}", level=2)
            # Confidence line
            p = doc.add_paragraph()
            run = p.add_run(f"Confidence: {r.confidence} ({r.confidence_pct}%)")
            run.font.color.rgb = (RGBColor(0,160,100) if r.confidence=="Yes"
                                  else RGBColor(200,0,0) if r.confidence=="No"
                                  else RGBColor(200,130,0))
            run.bold = True
            # Answer
            doc.add_paragraph(r.explanation)
            # Policy sources
            if r.document_names:
                doc.add_paragraph(f"Policy sources: {r.document_names}", style="Intense Quote")
            # Approval
            approver = get_approver(r)
            if approver:
                ap = doc.add_paragraph()
                ap.add_run(f"✓ Approved by: {approver}").font.color.rgb = RGBColor(0,160,100)
                ap.runs[0].bold = True
            doc.add_paragraph("")
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": "attachment; filename=SecureAnswer_Results.docx"})

    if fmt == "pdf":
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = [
            Paragraph("CybeSure SecureAnswer — Compliance Results", styles["Title"]),
            Paragraph(f"© {datetime.utcnow().year} CybeSure Ltd. All rights reserved. SecureAnswer™", styles["Normal"]),
            Spacer(1, 20)
        ]
        for r in results:
            story.append(Paragraph(f"<b>Q{r.question_id+1}:</b> {r.question}", styles["Heading2"]))
            conf_color = "green" if r.confidence=="Yes" else "red" if r.confidence=="No" else "orange"
            story.append(Paragraph(f"<b><font color='{conf_color}'>Confidence: {r.confidence} ({r.confidence_pct}%)</font></b>", styles["Normal"]))
            story.append(Paragraph(r.explanation, styles["Normal"]))
            if r.document_names:
                story.append(Paragraph(f"<i>Sources: {r.document_names}</i>", styles["Normal"]))
            approver = get_approver(r)
            if approver:
                story.append(Paragraph(f"<b><font color='green'>✓ Approved by: {approver}</font></b>", styles["Normal"]))
            story.append(Spacer(1, 12))
        doc.build(story)
        buf.seek(0)
        return StreamingResponse(buf, media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=SecureAnswer_Results.pdf"})
